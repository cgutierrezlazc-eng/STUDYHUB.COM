"""Tests del endpoint rúbrica (POST text, POST upload, GET) del 2d.6."""

from __future__ import annotations

import io
import os
import sys

import pytest

fastapi = pytest.importorskip("fastapi", reason="fastapi no instalado")
pytest.importorskip("httpx", reason="httpx no instalado")

from fastapi import FastAPI  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402
from sqlalchemy import create_engine, func  # noqa: E402
from sqlalchemy.orm import sessionmaker  # noqa: E402
from sqlalchemy.pool import StaticPool  # noqa: E402

_BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _BACKEND_DIR not in sys.path:
    sys.path.insert(0, _BACKEND_DIR)

os.environ.pop("DATABASE_URL", None)
os.environ.setdefault("JWT_SECRET", "test-secret-do-not-use-in-production")
os.environ.setdefault("OWNER_PASSWORD", "test-owner-password")
os.environ.setdefault("CORS_ORIGINS", "http://localhost:5173")
os.environ.setdefault("BCRYPT_ROUNDS", "4")


@pytest.fixture(scope="module")
def app_rubric():
    from database import Base, get_db  # type: ignore
    from workspaces_routes import router as ws_router  # type: ignore

    app = FastAPI()
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

    def override_get_db():
        db = TestingSessionLocal()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = override_get_db
    app.include_router(ws_router)
    yield app, TestingSessionLocal()


@pytest.fixture
def rubric_client(app_rubric):
    from datetime import datetime, timedelta

    from jose import jwt

    app, db = app_rubric
    from database import User, WorkspaceDocument, WorkspaceMember, gen_id  # type: ignore

    max_num = db.query(func.max(User.user_number)).scalar() or 0
    owner = User(
        id=gen_id(),
        email=f"rub_o_{gen_id()}@conniku.com",
        username=f"rub_o_{gen_id()}",
        user_number=max_num + 1,
        first_name="Rub",
        last_name="Owner",
        password_hash="$2b$04$test.hash.placeholder",
        birth_date="1995-05-15",
        created_at=datetime.utcnow(),
    )
    db.add(owner)
    db.flush()

    max_num = db.query(func.max(User.user_number)).scalar() or 0
    outsider = User(
        id=gen_id(),
        email=f"rub_out_{gen_id()}@conniku.com",
        username=f"rub_out_{gen_id()}",
        user_number=max_num + 1,
        first_name="Rub",
        last_name="Out",
        password_hash="$2b$04$test.hash.placeholder",
        birth_date="1995-05-15",
        created_at=datetime.utcnow(),
    )
    db.add(outsider)
    db.commit()
    db.refresh(owner)
    db.refresh(outsider)

    doc = WorkspaceDocument(
        id=gen_id(),
        title="Doc Rubric",
        owner_id=owner.id,
        apa_edition="7",
        options="{}",
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )
    db.add(doc)
    db.flush()
    db.add(
        WorkspaceMember(
            id=gen_id(),
            workspace_id=doc.id,
            user_id=owner.id,
            role="owner",
            invited_at=datetime.utcnow(),
            joined_at=datetime.utcnow(),
        )
    )
    db.commit()
    db.refresh(doc)

    secret = os.environ["JWT_SECRET"]
    exp = datetime.utcnow() + timedelta(hours=1)
    t_o = jwt.encode({"sub": owner.id, "exp": exp, "type": "access"}, secret, algorithm="HS256")
    t_out = jwt.encode({"sub": outsider.id, "exp": exp, "type": "access"}, secret, algorithm="HS256")

    with TestClient(app) as client:
        yield client, {"Authorization": f"Bearer {t_o}"}, {"Authorization": f"Bearer {t_out}"}, doc.id


# ─── Tests ────────────────────────────────────────────────────────


def test_rubric_text_ok(rubric_client):
    client, h_o, _, doc_id = rubric_client
    text = "Criterios:\n- Introducción (10 pts)\n- Desarrollo (30 pts)\n- Conclusión (20 pts)"
    r = client.post(f"/workspaces/{doc_id}/rubric/text", json={"text": text}, headers=h_o)
    assert r.status_code == 200
    body = r.json()
    assert len(body["items"]) == 3
    assert body["items"][0]["points"] == 10


def test_rubric_text_no_miembro_403(rubric_client):
    client, _, h_out, doc_id = rubric_client
    r = client.post(f"/workspaces/{doc_id}/rubric/text", json={"text": "- a (1 pts)"}, headers=h_out)
    assert r.status_code == 403


def test_rubric_text_vacio_400(rubric_client):
    client, h_o, _, doc_id = rubric_client
    r = client.post(f"/workspaces/{doc_id}/rubric/text", json={"text": "  "}, headers=h_o)
    assert r.status_code == 400


def test_rubric_get_despues_de_text(rubric_client):
    client, h_o, _, doc_id = rubric_client
    client.post(
        f"/workspaces/{doc_id}/rubric/text",
        json={"text": "- Item A (5 pts)\n- Item B (7 pts)"},
        headers=h_o,
    )
    r = client.get(f"/workspaces/{doc_id}/rubric", headers=h_o)
    assert r.status_code == 200
    body = r.json()
    assert len(body["items"]) == 2
    assert "Item A" in body["raw"]


def test_rubric_upload_txt_ok(rubric_client):
    client, h_o, _, doc_id = rubric_client
    content = b"Rubrica:\n- Ortografia (5 pts)\n- Claridad (10 pts)\n"
    r = client.post(
        f"/workspaces/{doc_id}/rubric/upload",
        files={"file": ("rubrica.txt", io.BytesIO(content), "text/plain")},
        headers=h_o,
    )
    assert r.status_code == 200
    body = r.json()
    assert len(body["items"]) == 2
    assert body["filename"] == "rubrica.txt"


def test_rubric_upload_sin_archivo_400(rubric_client):
    client, h_o, _, doc_id = rubric_client
    r = client.post(f"/workspaces/{doc_id}/rubric/upload", headers=h_o)
    assert r.status_code in (400, 422)


def test_rubric_upload_docx_real(rubric_client):
    """Sube un DOCX real generado en memoria con python-docx."""
    import docx as _docx

    client, h_o, _, doc_id = rubric_client

    buf = io.BytesIO()
    d = _docx.Document()
    d.add_paragraph("Criterios de evaluación:")
    d.add_paragraph("- Estructura (15 pts)")
    d.add_paragraph("- Contenido (20 pts)")
    d.save(buf)
    buf.seek(0)

    r = client.post(
        f"/workspaces/{doc_id}/rubric/upload",
        files={
            "file": (
                "rubrica.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=h_o,
    )
    assert r.status_code == 200
    body = r.json()
    assert len(body["items"]) == 2


def test_rubric_upload_excede_tamano_maximo_413(rubric_client):
    """Archivo > MAX_RUBRIC_UPLOAD_BYTES debe retornar 413."""
    client, h_o, _, doc_id = rubric_client
    oversized = b"A" * (11 * 1024 * 1024)  # 11 MB > 10 MB
    r = client.post(
        f"/workspaces/{doc_id}/rubric/upload",
        files={"file": ("gigante.txt", io.BytesIO(oversized), "text/plain")},
        headers=h_o,
    )
    assert r.status_code == 413, (
        f"Esperaba 413 Payload Too Large, got {r.status_code}: {r.text[:200]}"
    )


def test_rubric_upload_mime_no_permitido_415(rubric_client):
    """MIME fuera de allowlist (ej: image/png) debe retornar 415."""
    client, h_o, _, doc_id = rubric_client
    r = client.post(
        f"/workspaces/{doc_id}/rubric/upload",
        files={"file": ("pwn.png", io.BytesIO(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100), "image/png")},
        headers=h_o,
    )
    assert r.status_code == 415, (
        f"Esperaba 415 Unsupported Media Type, got {r.status_code}: {r.text[:200]}"
    )


def test_rubric_upload_extension_no_permitida_415(rubric_client):
    """Extensión fuera de allowlist (ej: .exe) debe retornar 415 incluso si MIME es generic."""
    client, h_o, _, doc_id = rubric_client
    r = client.post(
        f"/workspaces/{doc_id}/rubric/upload",
        files={
            "file": (
                "malware.exe",
                io.BytesIO(b"MZ\x90\x00" + b"\x00" * 100),
                "application/octet-stream",
            )
        },
        headers=h_o,
    )
    assert r.status_code == 415, (
        f"Esperaba 415 Unsupported Media Type, got {r.status_code}: {r.text[:200]}"
    )
