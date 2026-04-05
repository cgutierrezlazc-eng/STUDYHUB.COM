"""
Conniku Document Generator
===========================
Professional branded document generation for Conniku study platform.
Supports Word (.docx) and PDF export with consistent branding.

Brand:
  Primary: #2D62C8  |  Dark: #151B1E  |  Light BG: #F5F7F8
  Font: Calibri (Word) / Helvetica (PDF)
"""

import base64
import io
import os
import re
import tempfile
from datetime import datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# Conniku brand constants
# ---------------------------------------------------------------------------
BRAND_PRIMARY = (45, 98, 200)       # #2D62C8
BRAND_DARK = (21, 27, 30)           # #151B1E
BRAND_LIGHT_BG = (245, 247, 248)    # #F5F7F8
BRAND_GRAY = (128, 136, 146)
BRAND_WHITE = (255, 255, 255)
BRAND_PRIMARY_HEX = "#2D62C8"
BRAND_DARK_HEX = "#151B1E"

CONNIKU_DIR = Path.home() / ".conniku"
CONNIKU_DIR.mkdir(exist_ok=True)

# Embedded 80x80 PNG logo: blue square with white "C"
# (minimal hand-crafted PNG so the file stays self-contained)
_LOGO_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAFAAAABQCAYAAACOEfKtAAAACXBIWXMAAAsTAAAL"
    "EwEAmpwYAAADHUlEQVR4nO2dS07DMBCG/4EFO+AIsOIInIITcATOwBFYseAM"
    "HKGsWHIEjsAROAPskJBA4jXjxHHi+CfSqErtzHzxeOzYbgAAAAAAAAAAAAAA"
    "AAAAAIiJ1B0gInqQ9IakS0mXSLqQ1Ol0/vR6vY8oHWOm1+t96HQ6f0i6lHSJ"
    "pBuSHqR0Yneo2+1+yrL8K4Q4F0KcMcbOGGNnQohzKeVfIcR5t9v9jNKxFMjz"
    "/FsIcS6EOBNCnEkpz4QQZ1LKMyHEmRDiXAjxNcuyb1E6VgIp5V8hxLmU8kwI"
    "cSalPJNSnkkpz6WU50KI8yzLvkfpmBDigaRLSZdIuiTpUtIlki6RdImkSyRd"
    "Iukyk1JiSomRdInJZHIxnU4vpJSXUspLKeWllPJSCHEppbxM0xRTSqxKICNp"
    "QtKEMTZhjE0YYxPG2CRJknGapmOSxowxTCmx0sc4Hk8upJSXjLFLxtiEMTZh"
    "jE0YYxPG2IQxNmGMTdI0Hfv+h2OMsc/j8eRCSnnJGJswxiaMsQljbJIkyThN"
    "0zFjDNN1jJ1OZzQajS6EEOdSynMp5bkQ4lwIcS6EOE+S5Hw0Go1ardbfVqt1"
    "0Wq1RlE65nofc3X1+0xKeSaEOBdCnKdpejYcDs+GQ3E2HIqz4VCcDQbibDAQ"
    "Z/2+OOv3xVm/L876fXHW64mzXk+c9XrirNcTZ92uOOt2xVm3K866XXF2cSHO"
    "Li7E2cWFODs/F2fn5+Ls/Fyc9fvirNcTZ72eOOt2xVm3K866XXF2cSHOLi7E"
    "2fm5ODs/F2eDgTjr98VZvy/O+n1xNhyKs+FQnA2H4mwwEGd5Ls5Go3/GVfew"
    "Uq12kSTJOE3TMWNszBgbJ0kyzrJsTNI4SZIxY2ycpumYMTZO03TsS8dKoJTy"
    "Ukl5RdJVkiTjJEnGjLFxmqZjxtg4TdMxY2ycJMk4y7IxSeMkScb6WK4ElpR/"
    "kiQZSymvpJRXjLErxtiVlPKKpCuSrhhjV4yxqyRJxlmWjYUQ4yzLxlE65oph"
    "FY2J3SFjFGJ3iFHiOwAAAAAAAAAAAAAAAAAAAADQRP4DbLAHkbh4dCsAAAAASUVO"
    "RK5CYII="
)


def _get_logo_bytes() -> bytes:
    """Return the decoded logo PNG bytes."""
    return base64.b64decode(_LOGO_B64)


def _logo_stream() -> io.BytesIO:
    """Return a seekable stream of the logo PNG."""
    stream = io.BytesIO(_get_logo_bytes())
    stream.seek(0)
    return stream


def _temp_path(suffix: str) -> str:
    """Create a named temp file inside ~/.conniku/ and return its path."""
    fd, path = tempfile.mkstemp(suffix=suffix, dir=str(CONNIKU_DIR))
    os.close(fd)
    return path


def _today() -> str:
    return datetime.now().strftime("%d de %B de %Y")


# ===================================================================
#  DOCX helpers
# ===================================================================
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ===================================================================
#  PDF helpers  (reportlab)
# ===================================================================
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.units import inch, cm, mm
    from reportlab.lib.colors import HexColor, white, black, Color
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Frame, PageTemplate, BaseDocTemplate, Image,
        ListFlowable, ListItem, KeepTogether, HRFlowable,
    )
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.pdfgen import canvas as pdfcanvas
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


# ===================================================================
# BRAND COLORS for python-docx
# ===================================================================
def _rgb_primary():
    return RGBColor(*BRAND_PRIMARY)

def _rgb_dark():
    return RGBColor(*BRAND_DARK)

def _rgb_gray():
    return RGBColor(*BRAND_GRAY)

def _rgb_white():
    return RGBColor(*BRAND_WHITE)


# ===================================================================
#  DOCX: internal styling helpers
# ===================================================================

def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_shading(paragraph, color_hex: str):
    """Apply background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    pPr.append(shading)


def _add_formatted_runs(paragraph, text: str, default_size=Pt(11), default_font="Calibri"):
    """Parse **bold**, *italic*, ~~strikethrough~~, `code`, and $latex$ into Word runs."""
    # Order matters: bold first, then italic, then strike, then code, then latex
    pattern = r'(\*\*.*?\*\*|~~.*?~~|\*.*?\*|`[^`]+`|\$[^$]+\$)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("~~") and part.endswith("~~"):
            run = paragraph.add_run(part[2:-2])
            run.font.strike = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        elif part.startswith("$") and part.endswith("$"):
            run = paragraph.add_run(part)  # keep $ delimiters so formulas are visible
            run.italic = True
            run.font.name = "Cambria Math"
            run.font.size = default_size
        else:
            run = paragraph.add_run(part)
        run.font.size = run.font.size or default_size
        run.font.name = run.font.name or default_font


def _configure_styles(doc):
    """Set up document-wide styles for the Conniku brand."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = _rgb_dark()
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.color.rgb = _rgb_primary()
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        heading_style.paragraph_format.space_after = Pt(6)
        sizes = {1: Pt(22), 2: Pt(17), 3: Pt(14)}
        heading_style.font.size = sizes[level]


def _add_cover_page(doc, title: str, subtitle: str = "Material de Estudio"):
    """Add a branded cover page."""
    # Blue header bar as a full-width table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(2)
    tbl.columns[1].width = Inches(4.5)

    cell_logo = tbl.cell(0, 0)
    cell_title = tbl.cell(0, 1)
    _set_cell_shading(cell_logo, "2D62C8")
    _set_cell_shading(cell_title, "2D62C8")

    # Logo in left cell
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p_logo.add_run().add_picture(_logo_stream(), width=Inches(0.7))
    except Exception:
        run = p_logo.add_run("C")
        run.font.size = Pt(36)
        run.font.color.rgb = _rgb_white()
        run.bold = True

    # Title in right cell
    p_title = cell_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_t = p_title.add_run(title)
    run_t.font.size = Pt(28)
    run_t.font.color.rgb = _rgb_white()
    run_t.bold = True
    run_t.font.name = "Calibri"

    p_sub = cell_title.add_paragraph()
    run_s = p_sub.add_run(subtitle)
    run_s.font.size = Pt(14)
    run_s.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    run_s.font.name = "Calibri"

    # Spacer
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Date & attribution
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = p_date.add_run(_today())
    run_d.font.size = Pt(12)
    run_d.font.color.rgb = _rgb_gray()

    doc.add_paragraph("")

    p_attr = doc.add_paragraph()
    p_attr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = p_attr.add_run("Generado con Conniku \u2014 conniku.com")
    run_a.font.size = Pt(10)
    run_a.font.color.rgb = _rgb_primary()
    run_a.italic = True

    # Page break after cover
    doc.add_page_break()


def _add_header_footer(doc):
    """Add header (logo + Conniku) and footer (copyright + page number) to every page."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    h_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_table.autofit = False
    h_table.columns[0].width = Inches(4)
    h_table.columns[1].width = Inches(2.5)

    # Left: logo + text
    left_cell = h_table.cell(0, 0)
    p_left = left_cell.paragraphs[0]
    try:
        p_left.add_run().add_picture(_logo_stream(), width=Inches(0.25))
    except Exception:
        pass
    run_name = p_left.add_run("  Conniku")
    run_name.font.size = Pt(9)
    run_name.font.color.rgb = _rgb_primary()
    run_name.bold = True
    run_name.font.name = "Calibri"

    # Right: page number placeholder
    right_cell = h_table.cell(0, 1)
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_pg = p_right.add_run()
    run_pg.font.size = Pt(9)
    run_pg.font.color.rgb = _rgb_gray()
    # Insert PAGE field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_pg._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_pg._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_pg._r.append(fldChar2)

    # Remove default header paragraph border
    for p in header.paragraphs:
        if not p.text.strip():
            p.clear()

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run("\u00A9 Conniku \u2014 Material de estudio personal")
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = _rgb_gray()
    run_f.font.name = "Calibri"


def _add_toc(doc):
    """Insert an auto-updating Table of Contents."""
    p_title = doc.add_paragraph()
    run = p_title.add_run("Tabla de Contenidos")
    run.font.size = Pt(18)
    run.font.color.rgb = _rgb_primary()
    run.bold = True
    run.font.name = "Calibri"

    # Insert TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._r.append(fldChar2)
    run_placeholder = paragraph.add_run("(Actualizar tabla de contenidos: clic derecho > Actualizar campo)")
    run_placeholder.font.size = Pt(9)
    run_placeholder.font.color.rgb = _rgb_gray()
    run_placeholder.italic = True
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_placeholder._r.append(fldChar3)

    doc.add_paragraph("")
    doc.add_page_break()


def _parse_table_block(lines: list):
    """Parse a markdown table block into list of rows (list of cell strings)."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            # skip separator rows
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            rows.append(cells)
    return rows


def _add_docx_table(doc, rows):
    """Add a styled table to the document."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < n_cols:
                cell = tbl.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_formatted_runs(p, cell_text, default_size=Pt(10))
                if i == 0:
                    _set_cell_shading(cell, "2D62C8")
                    for run in p.runs:
                        run.font.color.rgb = _rgb_white()
                        run.bold = True
    doc.add_paragraph("")


# ===================================================================
#  PUBLIC: markdown_to_docx
# ===================================================================

def markdown_to_docx(content: str, title: str = "Conniku Document") -> str:
    """
    Convert markdown text to a professionally branded .docx file.
    Returns the path to the generated file in ~/.conniku/.
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()
    _configure_styles(doc)
    _add_header_footer(doc)
    _add_cover_page(doc, title)
    _add_toc(doc)

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""
    table_lines = []
    in_table = False
    in_key_box = False
    key_box_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Key term boxes: :::key ... ::: ---
        if stripped == ":::key":
            in_key_box = True
            key_box_lines = []
            i += 1
            continue
        if stripped == ":::" and in_key_box:
            in_key_box = False
            # Render key box as shaded table
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            _set_cell_shading(cell, "E8EFF9")
            for kl in key_box_lines:
                p = cell.paragraphs[0] if cell.paragraphs[0].text == "" and key_box_lines.index(kl) == 0 else cell.add_paragraph()
                _add_formatted_runs(p, kl.strip(), default_size=Pt(10))
            # Blue left border
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="24" w:space="0" w:color="2D62C8"/>'
                '</w:tcBorders>'
            )
            tc_pr.append(borders)
            doc.add_paragraph("")
            i += 1
            continue
        if in_key_box:
            key_box_lines.append(line)
            i += 1
            continue

        # --- Code blocks ---
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                _set_cell_shading(cell, "F0F0F0")
                p = cell.paragraphs[0]
                if code_lang:
                    run_lang = p.add_run(code_lang.upper() + "\n")
                    run_lang.font.size = Pt(7)
                    run_lang.font.color.rgb = _rgb_gray()
                    run_lang.bold = True
                run_code = p.add_run(code_text)
                run_code.font.name = "Courier New"
                run_code.font.size = Pt(9)
                run_code.font.color.rgb = _rgb_dark()
                code_lines = []
                code_lang = ""
                in_code_block = False
                doc.add_paragraph("")
            else:
                in_code_block = True
                code_lang = stripped[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Table detection ---
        if stripped.startswith("|") and stripped.endswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            in_table = False
            rows = _parse_table_block(table_lines)
            _add_docx_table(doc, rows)
            table_lines = []
            # Don't increment i — process current line normally
            continue

        # --- Horizontal rule ---
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            border = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                '</w:pBdr>'
            )
            p_pr.append(border)
            i += 1
            continue

        # --- Headings ---
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # --- Blockquotes ---
        if stripped.startswith("> "):
            quote_text = stripped[2:]
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            _set_cell_shading(cell, "F5F7F8")
            p = cell.paragraphs[0]
            _add_formatted_runs(p, quote_text, default_size=Pt(11))
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            # Blue left border
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="24" w:space="0" w:color="2D62C8"/>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tc_pr.append(borders)
            i += 1
            continue

        # --- Bullet lists ---
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        # --- Numbered lists ---
        m = re.match(r"^(\d+)\.\s(.+)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, m.group(2))
            i += 1
            continue

        # --- Empty lines ---
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        _add_formatted_runs(p, stripped)
        i += 1

    # Flush pending table
    if in_table and table_lines:
        rows = _parse_table_block(table_lines)
        _add_docx_table(doc, rows)

    # Save
    path = _temp_path(".docx")
    doc.save(path)
    return path


# ===================================================================
#  PUBLIC: summary_to_docx
# ===================================================================

def summary_to_docx(summary_data: dict, title: str = "Resumen de Estudio") -> str:
    """
    Generate a branded study summary document from structured AI data.

    Expected summary_data keys:
      - sections: list of {title, content, keyPoints?}
      - keyTerms: list of {term, definition}
      - formulas: list of {name, formula, explanation}
      - studyTips: list of str

    Returns the path to the generated .docx file.
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()
    _configure_styles(doc)
    _add_header_footer(doc)
    _add_cover_page(doc, title, "Resumen de Estudio")
    _add_toc(doc)

    sections = summary_data.get("sections", [])
    key_terms = summary_data.get("keyTerms", [])
    formulas = summary_data.get("formulas", [])
    study_tips = summary_data.get("studyTips", [])

    # --- Sections ---
    for idx, section in enumerate(sections):
        doc.add_heading(section.get("title", f"Secci\u00F3n {idx+1}"), level=1)
        content = section.get("content", "")
        if content:
            for para_text in content.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph()
                    _add_formatted_runs(p, para_text)

        key_points = section.get("keyPoints", [])
        if key_points:
            doc.add_paragraph("")
            p_kp_title = doc.add_paragraph()
            run_kp = p_kp_title.add_run("Puntos Clave")
            run_kp.bold = True
            run_kp.font.color.rgb = _rgb_primary()
            run_kp.font.size = Pt(12)
            for kp in key_points:
                p = doc.add_paragraph(style="List Bullet")
                _add_formatted_runs(p, kp)

        if idx < len(sections) - 1:
            doc.add_paragraph("")

    # --- Key Terms (Glossary Table) ---
    if key_terms:
        doc.add_page_break()
        doc.add_heading("Glosario de T\u00E9rminos Clave", level=1)
        doc.add_paragraph("")

        tbl = doc.add_table(rows=1 + len(key_terms), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        tbl.columns[0].width = Inches(2.5)
        tbl.columns[1].width = Inches(4)

        # Header row
        for j, header_text in enumerate(["T\u00E9rmino", "Definici\u00F3n"]):
            cell = tbl.cell(0, j)
            _set_cell_shading(cell, "2D62C8")
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.bold = True
            run.font.color.rgb = _rgb_white()
            run.font.size = Pt(11)

        for i, kt in enumerate(key_terms):
            term_cell = tbl.cell(i + 1, 0)
            def_cell = tbl.cell(i + 1, 1)
            p_term = term_cell.paragraphs[0]
            run_t = p_term.add_run(kt.get("term", ""))
            run_t.bold = True
            run_t.font.size = Pt(10)

            p_def = def_cell.paragraphs[0]
            _add_formatted_runs(p_def, kt.get("definition", ""), default_size=Pt(10))

            # Zebra striping
            if i % 2 == 0:
                _set_cell_shading(term_cell, "F5F7F8")
                _set_cell_shading(def_cell, "F5F7F8")

        doc.add_paragraph("")

    # --- Formulas ---
    if formulas:
        doc.add_page_break()
        doc.add_heading("F\u00F3rmulas", level=1)
        doc.add_paragraph("")

        tbl = doc.add_table(rows=1 + len(formulas), cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        tbl.columns[0].width = Inches(1.8)
        tbl.columns[1].width = Inches(2.2)
        tbl.columns[2].width = Inches(2.5)

        for j, header_text in enumerate(["Nombre", "F\u00F3rmula", "Explicaci\u00F3n"]):
            cell = tbl.cell(0, j)
            _set_cell_shading(cell, "2D62C8")
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.bold = True
            run.font.color.rgb = _rgb_white()
            run.font.size = Pt(11)

        for i, fm in enumerate(formulas):
            tbl.cell(i + 1, 0).paragraphs[0].add_run(fm.get("name", "")).font.size = Pt(10)
            formula_run = tbl.cell(i + 1, 1).paragraphs[0].add_run(fm.get("formula", ""))
            formula_run.font.name = "Cambria Math"
            formula_run.italic = True
            formula_run.font.size = Pt(10)
            _add_formatted_runs(
                tbl.cell(i + 1, 2).paragraphs[0],
                fm.get("explanation", ""),
                default_size=Pt(10),
            )
            if i % 2 == 0:
                for j in range(3):
                    _set_cell_shading(tbl.cell(i + 1, j), "F5F7F8")

        doc.add_paragraph("")

    # --- Study Tips ---
    if study_tips:
        doc.add_page_break()
        doc.add_heading("Consejos de Estudio", level=1)
        doc.add_paragraph("")

        for tip in study_tips:
            p = doc.add_paragraph()
            # Checkbox character
            run_cb = p.add_run("\u2610  ")
            run_cb.font.size = Pt(13)
            _add_formatted_runs(p, tip)

    # --- Notas personales ---
    doc.add_page_break()
    doc.add_heading("Notas Personales", level=1)
    p_note_intro = doc.add_paragraph()
    run_ni = p_note_intro.add_run("Utiliza este espacio para escribir tus propias notas y observaciones:")
    run_ni.italic = True
    run_ni.font.color.rgb = _rgb_gray()

    # Add ruled lines
    for _ in range(20):
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        border = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="DDDDDD"/>'
            '</w:pBdr>'
        )
        p_pr.append(border)
        p.paragraph_format.space_after = Pt(12)

    path = _temp_path(".docx")
    doc.save(path)
    return path


# ===================================================================
#  PDF HELPERS
# ===================================================================

def _pdf_styles():
    """Build a dictionary of branded ParagraphStyles for reportlab."""
    styles = getSampleStyleSheet()

    brand_blue = HexColor(BRAND_PRIMARY_HEX)
    brand_dark = HexColor(BRAND_DARK_HEX)

    custom = {}

    custom["CoverTitle"] = ParagraphStyle(
        "CoverTitle", parent=styles["Title"],
        fontName="Helvetica-Bold", fontSize=28, leading=34,
        textColor=white, alignment=TA_LEFT,
    )
    custom["CoverSubtitle"] = ParagraphStyle(
        "CoverSubtitle", parent=styles["Normal"],
        fontName="Helvetica", fontSize=14, leading=18,
        textColor=HexColor("#CCDDFF"), alignment=TA_LEFT,
    )
    custom["CoverDate"] = ParagraphStyle(
        "CoverDate", parent=styles["Normal"],
        fontName="Helvetica", fontSize=12, leading=16,
        textColor=HexColor("#888888"), alignment=TA_CENTER,
    )
    custom["CoverAttrib"] = ParagraphStyle(
        "CoverAttrib", parent=styles["Normal"],
        fontName="Helvetica-Oblique", fontSize=10, leading=14,
        textColor=brand_blue, alignment=TA_CENTER,
    )
    custom["H1"] = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=20, leading=26,
        textColor=brand_blue, spaceAfter=8, spaceBefore=18,
    )
    custom["H2"] = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontName="Helvetica-Bold", fontSize=16, leading=20,
        textColor=brand_blue, spaceAfter=6, spaceBefore=14,
    )
    custom["H3"] = ParagraphStyle(
        "H3", parent=styles["Heading3"],
        fontName="Helvetica-Bold", fontSize=13, leading=17,
        textColor=brand_blue, spaceAfter=4, spaceBefore=10,
    )
    custom["Body"] = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10.5, leading=15,
        textColor=brand_dark, spaceAfter=6, alignment=TA_JUSTIFY,
    )
    custom["Bullet"] = ParagraphStyle(
        "Bullet", parent=custom["Body"],
        leftIndent=20, bulletIndent=8,
        spaceAfter=3,
    )
    custom["Code"] = ParagraphStyle(
        "Code", parent=styles["Code"],
        fontName="Courier", fontSize=9, leading=12,
        textColor=brand_dark, backColor=HexColor("#F0F0F0"),
        leftIndent=12, rightIndent=12,
        spaceAfter=6, spaceBefore=6,
    )
    custom["Blockquote"] = ParagraphStyle(
        "Blockquote", parent=custom["Body"],
        fontName="Helvetica-Oblique", fontSize=10.5,
        textColor=HexColor("#555555"),
        leftIndent=18, borderPadding=6,
    )
    custom["Footer"] = ParagraphStyle(
        "Footer", parent=styles["Normal"],
        fontName="Helvetica", fontSize=7.5, leading=10,
        textColor=HexColor("#888888"), alignment=TA_CENTER,
    )
    custom["TOCEntry"] = ParagraphStyle(
        "TOCEntry", parent=custom["Body"],
        fontSize=11, leading=18, leftIndent=10,
    )
    custom["ChatStudent"] = ParagraphStyle(
        "ChatStudent", parent=custom["Body"],
        fontName="Helvetica-Bold", fontSize=10.5,
        textColor=brand_dark, leftIndent=0,
    )
    custom["ChatConniku"] = ParagraphStyle(
        "ChatConniku", parent=custom["Body"],
        fontName="Helvetica", fontSize=10.5,
        textColor=brand_blue, leftIndent=0,
    )
    custom["ChatMeta"] = ParagraphStyle(
        "ChatMeta", parent=styles["Normal"],
        fontName="Helvetica", fontSize=8, leading=10,
        textColor=HexColor("#999999"),
    )
    custom["Checkbox"] = ParagraphStyle(
        "Checkbox", parent=custom["Body"],
        leftIndent=10, spaceAfter=4,
    )

    return custom


def _pdf_header_footer(canvas_obj, doc):
    """Draw header and footer on every PDF page."""
    canvas_obj.saveState()
    width, height = A4

    # Header line
    canvas_obj.setStrokeColor(HexColor(BRAND_PRIMARY_HEX))
    canvas_obj.setLineWidth(1.5)
    canvas_obj.line(50, height - 40, width - 50, height - 40)

    # Header text
    canvas_obj.setFont("Helvetica-Bold", 8)
    canvas_obj.setFillColor(HexColor(BRAND_PRIMARY_HEX))
    canvas_obj.drawString(55, height - 35, "Conniku")

    # Page number
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(HexColor("#888888"))
    canvas_obj.drawRightString(width - 55, height - 35, f"P\u00E1gina {doc.page}")

    # Footer
    canvas_obj.setFont("Helvetica", 7)
    canvas_obj.setFillColor(HexColor("#888888"))
    canvas_obj.drawCentredString(
        width / 2, 25,
        "\u00A9 Conniku \u2014 Material de estudio personal"
    )
    canvas_obj.restoreState()


def _pdf_cover_page(story, styles, title: str, subtitle: str = "Material de Estudio"):
    """Build the cover page flowables for a PDF."""
    width, _ = A4
    content_width = width - 100  # margins

    # Blue header bar
    bar_data = [[""]]
    bar = Table(bar_data, colWidths=[content_width], rowHeights=[120])
    bar.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor(BRAND_PRIMARY_HEX)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))

    # Title inside blue bar — use a nested table
    title_para = Paragraph(title, styles["CoverTitle"])
    sub_para = Paragraph(subtitle, styles["CoverSubtitle"])
    inner_data = [[title_para], [sub_para]]
    inner_table = Table(inner_data, colWidths=[content_width - 40])
    inner_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor(BRAND_PRIMARY_HEX)),
        ("LEFTPADDING", (0, 0), (-1, -1), 20),
        ("TOPPADDING", (0, 0), (0, 0), 20),
        ("BOTTOMPADDING", (-1, -1), (-1, -1), 20),
    ]))

    story.append(inner_table)
    story.append(Spacer(1, 60))
    story.append(Paragraph(_today(), styles["CoverDate"]))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Generado con Conniku \u2014 conniku.com", styles["CoverAttrib"]))
    story.append(PageBreak())


def _pdf_toc_page(story, styles, headings: list):
    """Build a simple table of contents from a list of (level, title) tuples."""
    story.append(Paragraph("Tabla de Contenidos", styles["H1"]))
    story.append(Spacer(1, 12))
    for level, heading_text in headings:
        indent = (level - 1) * 20
        s = ParagraphStyle(
            f"TOC{level}", parent=styles["TOCEntry"],
            leftIndent=indent,
            fontName="Helvetica-Bold" if level == 1 else "Helvetica",
            fontSize=11 if level == 1 else 10,
        )
        story.append(Paragraph(heading_text, s))
    story.append(PageBreak())


def _md_to_pdf_flowables(content: str, styles: dict) -> tuple:
    """Parse markdown into reportlab flowables. Returns (flowables, headings)."""
    flowables = []
    headings = []  # (level, text) for TOC
    lines = content.split("\n")
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code:
                code_text = "\n".join(code_lines).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                flowables.append(Paragraph(f"<pre>{code_text}</pre>", styles["Code"]))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            txt = _escape_html(stripped[4:])
            headings.append((3, txt))
            flowables.append(Paragraph(txt, styles["H3"]))
            i += 1
            continue
        if stripped.startswith("## "):
            txt = _escape_html(stripped[3:])
            headings.append((2, txt))
            flowables.append(Paragraph(txt, styles["H2"]))
            i += 1
            continue
        if stripped.startswith("# "):
            txt = _escape_html(stripped[2:])
            headings.append((1, txt))
            flowables.append(Paragraph(txt, styles["H1"]))
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            flowables.append(HRFlowable(
                width="100%", thickness=1, color=HexColor("#CCCCCC"),
                spaceAfter=8, spaceBefore=8,
            ))
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            txt = _md_inline_to_html(stripped[2:])
            # Render as indented italic with blue bar via table
            q_para = Paragraph(txt, styles["Blockquote"])
            bar_tbl = Table(
                [[q_para]],
                colWidths=[430],
            )
            bar_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), HexColor("#F5F7F8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LINEBEFOREDECOR", (0, 0), (0, -1), 3, HexColor(BRAND_PRIMARY_HEX)),
            ]))
            flowables.append(bar_tbl)
            flowables.append(Spacer(1, 4))
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            txt = _md_inline_to_html(stripped[2:])
            flowables.append(Paragraph(f"\u2022  {txt}", styles["Bullet"]))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s(.+)", stripped)
        if m:
            txt = _md_inline_to_html(m.group(2))
            flowables.append(Paragraph(f"{m.group(1)}.  {txt}", styles["Bullet"]))
            i += 1
            continue

        # Table
        if stripped.startswith("|") and stripped.endswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            rows = _parse_table_block(tbl_lines)
            if rows:
                table_data = []
                for ri, row in enumerate(rows):
                    table_data.append([
                        Paragraph(
                            _md_inline_to_html(cell),
                            ParagraphStyle(
                                "tc", fontName="Helvetica-Bold" if ri == 0 else "Helvetica",
                                fontSize=9, leading=12,
                                textColor=white if ri == 0 else HexColor(BRAND_DARK_HEX),
                            ),
                        )
                        for cell in row
                    ])
                n_cols = max(len(r) for r in table_data)
                col_w = 450 / n_cols
                t = Table(table_data, colWidths=[col_w] * n_cols)
                style_cmds = [
                    ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor(BRAND_PRIMARY_HEX)),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
                # Zebra stripe
                for ri in range(1, len(table_data)):
                    if ri % 2 == 0:
                        style_cmds.append(("BACKGROUND", (0, ri), (-1, ri), HexColor("#F5F7F8")))
                t.setStyle(TableStyle(style_cmds))
                flowables.append(t)
                flowables.append(Spacer(1, 8))
            continue

        # Empty line
        if not stripped:
            flowables.append(Spacer(1, 6))
            i += 1
            continue

        # Regular paragraph
        txt = _md_inline_to_html(stripped)
        flowables.append(Paragraph(txt, styles["Body"]))
        i += 1

    return flowables, headings


def _escape_html(text: str) -> str:
    """Escape HTML special chars."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _md_inline_to_html(text: str) -> str:
    """Convert inline markdown (bold, italic, code, strikethrough) to reportlab-compatible HTML."""
    text = _escape_html(text)
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    # Strikethrough
    text = re.sub(r'~~(.+?)~~', r'<strike>\1</strike>', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    # Inline code
    text = re.sub(r'`([^`]+)`', r'<font face="Courier" size="9">\1</font>', text)
    return text


def _build_pdf(story_builder_fn, title: str, subtitle: str = "Material de Estudio") -> str:
    """
    Generic PDF builder. story_builder_fn receives (story, styles) and populates the story.
    Returns temp file path.
    """
    if not HAS_PDF:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab")

    path = _temp_path(".pdf")
    styles = _pdf_styles()

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        topMargin=50,
        bottomMargin=45,
        leftMargin=50,
        rightMargin=50,
        title=title,
        author="Conniku",
    )

    story = []
    _pdf_cover_page(story, styles, title, subtitle)
    story_builder_fn(story, styles)

    doc.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
    return path


# ===================================================================
#  PUBLIC: summary_to_pdf
# ===================================================================

def summary_to_pdf(summary_data: dict, title: str = "Resumen de Estudio") -> str:
    """
    Generate a branded PDF study summary from structured AI data.

    Expected summary_data keys:
      - sections: list of {title, content, keyPoints?}
      - keyTerms: list of {term, definition}
      - formulas: list of {name, formula, explanation}
      - studyTips: list of str

    Returns the path to the generated .pdf file.
    """

    def build(story, styles):
        sections = summary_data.get("sections", [])
        key_terms = summary_data.get("keyTerms", [])
        formulas = summary_data.get("formulas", [])
        study_tips = summary_data.get("studyTips", [])

        # Collect headings for TOC
        headings = []
        for section in sections:
            headings.append((1, section.get("title", "Secci\u00F3n")))
        if key_terms:
            headings.append((1, "Glosario de T\u00E9rminos Clave"))
        if formulas:
            headings.append((1, "F\u00F3rmulas"))
        if study_tips:
            headings.append((1, "Consejos de Estudio"))
        headings.append((1, "Notas Personales"))

        _pdf_toc_page(story, styles, headings)

        # --- Sections ---
        for section in sections:
            story.append(Paragraph(_escape_html(section.get("title", "")), styles["H1"]))
            content = section.get("content", "")
            if content:
                for para in content.split("\n"):
                    para = para.strip()
                    if para:
                        story.append(Paragraph(_md_inline_to_html(para), styles["Body"]))

            key_points = section.get("keyPoints", [])
            if key_points:
                story.append(Spacer(1, 6))
                story.append(Paragraph(
                    "<b>Puntos Clave</b>",
                    ParagraphStyle("KP", parent=styles["Body"], textColor=HexColor(BRAND_PRIMARY_HEX), fontSize=12),
                ))
                for kp in key_points:
                    story.append(Paragraph(f"\u2022  {_md_inline_to_html(kp)}", styles["Bullet"]))

            story.append(Spacer(1, 12))

        # --- Key Terms ---
        if key_terms:
            story.append(PageBreak())
            story.append(Paragraph("Glosario de T\u00E9rminos Clave", styles["H1"]))
            story.append(Spacer(1, 8))

            header_style = ParagraphStyle(
                "KTHeader", fontName="Helvetica-Bold", fontSize=10, leading=14, textColor=white,
            )
            body_style = ParagraphStyle(
                "KTBody", fontName="Helvetica", fontSize=9.5, leading=13, textColor=HexColor(BRAND_DARK_HEX),
            )
            term_style = ParagraphStyle(
                "KTTerm", fontName="Helvetica-Bold", fontSize=9.5, leading=13, textColor=HexColor(BRAND_DARK_HEX),
            )

            table_data = [
                [Paragraph("T\u00E9rmino", header_style), Paragraph("Definici\u00F3n", header_style)]
            ]
            for kt in key_terms:
                table_data.append([
                    Paragraph(_escape_html(kt.get("term", "")), term_style),
                    Paragraph(_md_inline_to_html(kt.get("definition", "")), body_style),
                ])

            t = Table(table_data, colWidths=[150, 310])
            style_cmds = [
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                ("BACKGROUND", (0, 0), (-1, 0), HexColor(BRAND_PRIMARY_HEX)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
            for ri in range(1, len(table_data)):
                if ri % 2 == 0:
                    style_cmds.append(("BACKGROUND", (0, ri), (-1, ri), HexColor("#F5F7F8")))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)

        # --- Formulas ---
        if formulas:
            story.append(PageBreak())
            story.append(Paragraph("F\u00F3rmulas", styles["H1"]))
            story.append(Spacer(1, 8))

            header_style = ParagraphStyle(
                "FHeader", fontName="Helvetica-Bold", fontSize=10, leading=14, textColor=white,
            )
            body_style = ParagraphStyle(
                "FBody", fontName="Helvetica", fontSize=9.5, leading=13, textColor=HexColor(BRAND_DARK_HEX),
            )
            formula_style = ParagraphStyle(
                "FFormula", fontName="Courier", fontSize=10, leading=13, textColor=HexColor(BRAND_DARK_HEX),
            )

            table_data = [[
                Paragraph("Nombre", header_style),
                Paragraph("F\u00F3rmula", header_style),
                Paragraph("Explicaci\u00F3n", header_style),
            ]]
            for fm in formulas:
                table_data.append([
                    Paragraph(_escape_html(fm.get("name", "")), body_style),
                    Paragraph(_escape_html(fm.get("formula", "")), formula_style),
                    Paragraph(_md_inline_to_html(fm.get("explanation", "")), body_style),
                ])

            t = Table(table_data, colWidths=[120, 160, 180])
            style_cmds = [
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                ("BACKGROUND", (0, 0), (-1, 0), HexColor(BRAND_PRIMARY_HEX)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
            for ri in range(1, len(table_data)):
                if ri % 2 == 0:
                    style_cmds.append(("BACKGROUND", (0, ri), (-1, ri), HexColor("#F5F7F8")))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)

        # --- Study Tips ---
        if study_tips:
            story.append(PageBreak())
            story.append(Paragraph("Consejos de Estudio", styles["H1"]))
            story.append(Spacer(1, 8))
            for tip in study_tips:
                story.append(Paragraph(
                    f"\u2610  {_md_inline_to_html(tip)}",
                    styles["Checkbox"],
                ))

        # --- Notas personales ---
        story.append(PageBreak())
        story.append(Paragraph("Notas Personales", styles["H1"]))
        story.append(Spacer(1, 4))
        story.append(Paragraph(
            "<i>Utiliza este espacio para escribir tus propias notas y observaciones:</i>",
            ParagraphStyle("NI", parent=styles["Body"], textColor=HexColor("#888888")),
        ))
        story.append(Spacer(1, 12))
        # Ruled lines
        for _ in range(22):
            story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#DDDDDD"), spaceAfter=18))

    return _build_pdf(build, title, "Resumen de Estudio")


# ===================================================================
#  PUBLIC: chat_to_pdf
# ===================================================================

def chat_to_pdf(messages: list, title: str = "Chat de Estudio") -> str:
    """
    Export a chat conversation to a branded PDF.

    messages: list of dicts with keys:
      - role: "user" | "assistant" | "student" | "conniku"
      - content: str (may contain markdown)
      - timestamp: str (optional, ISO or display format)

    Returns the path to the generated .pdf file.
    """

    def build(story, styles):
        # Simple TOC
        story.append(Paragraph("Conversaci\u00F3n de Estudio", styles["H1"]))
        story.append(Spacer(1, 4))
        story.append(Paragraph(
            f"Total de mensajes: {len(messages)}",
            ParagraphStyle("MsgCount", parent=styles["Body"], textColor=HexColor("#888888")),
        ))
        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor(BRAND_PRIMARY_HEX), spaceAfter=12))

        for msg in messages:
            role_raw = msg.get("role", "user").lower()
            if role_raw in ("user", "student", "human"):
                role_label = "Estudiante"
                role_style = styles["ChatStudent"]
                bubble_bg = HexColor("#F5F7F8")
            else:
                role_label = "Conniku"
                role_style = styles["ChatConniku"]
                bubble_bg = HexColor("#E8EFF9")

            timestamp = msg.get("timestamp", "")
            content = msg.get("content", "")

            # Role + timestamp header
            meta_text = f"<b>{_escape_html(role_label)}</b>"
            if timestamp:
                meta_text += f"  &nbsp; <font color='#999999' size='8'>{_escape_html(str(timestamp))}</font>"
            story.append(Paragraph(meta_text, role_style))

            # Message content as a bubble
            content_html = _md_inline_to_html(content).replace("\n", "<br/>")
            body_para = Paragraph(content_html, styles["Body"])
            bubble = Table(
                [[body_para]],
                colWidths=[440],
            )
            bubble.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), bubble_bg),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("ROUNDEDCORNERS", [6, 6, 6, 6]),
            ]))
            story.append(bubble)
            story.append(Spacer(1, 10))

    return _build_pdf(build, title, "Chat de Estudio")
