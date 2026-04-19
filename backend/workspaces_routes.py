"""
Workspaces v2 — Redacción Colaborativa de Documentos Académicos.

Endpoints REST del módulo Workspaces. Prefijo /workspaces.
En sub-bloque 2a: CRUD de workspaces, miembros y versiones.
En sub-bloque 2b: chat grupal (GET/POST/DELETE /chat/messages),
    contribution metric (PATCH /members/{id}/contribution),
    content_yjs en PatchWorkspaceRequest.
No incluye: Yjs WS (2b, en workspaces_ws.py), Athena (2c),
    export/rúbrica/math/share (2d).

Patrón: sigue el estilo de collab_routes.py con _check_access helper.
"""

from __future__ import annotations

import logging
from datetime import datetime
from typing import Optional

from database import (
    User,
    WorkspaceDocument,
    WorkspaceMember,
    WorkspaceMessage,
    WorkspaceVersion,
    gen_id,
    get_db,
)
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from middleware import get_current_user
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

logger = logging.getLogger("conniku.workspaces")

router = APIRouter(prefix="/workspaces", tags=["workspaces"])

# ─── Hardening rubric upload (A-2, 2026-04-19) ────────────────────
MAX_RUBRIC_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB

ALLOWED_RUBRIC_MIME_TYPES = frozenset(
    {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "text/plain",
    }
)

ALLOWED_RUBRIC_EXTENSIONS = frozenset({".pdf", ".docx", ".doc", ".txt"})


# ─── Helpers internos ─────────────────────────────────────────────


def _user_brief(u: Optional[User]) -> Optional[dict]:
    """Devuelve representación mínima de un usuario para respuestas JSON."""
    if not u:
        return None
    return {
        "id": u.id,
        "username": u.username,
        "firstName": u.first_name,
        "lastName": u.last_name,
        "avatar": u.avatar,
    }


def _check_access(
    doc_id: str,
    user: User,
    db: Session,
    required_role: str = "viewer",
) -> tuple:
    """Verifica que el usuario tenga acceso al workspace.

    Devuelve (doc, member_record | None).
    Lanza HTTPException 404 si no existe, 403 si no tiene el rol requerido.

    Jerarquía de roles: viewer (0) < editor (1) < owner (2).
    - required_role='viewer': cualquier miembro puede acceder (lectura).
    - required_role='editor': editors y owner pueden modificar contenido.
    - required_role='owner': solo el owner (eliminar, gestionar roles).
    """
    doc = db.query(WorkspaceDocument).filter(WorkspaceDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Workspace no encontrado")

    # Owner siempre tiene acceso total
    if doc.owner_id == user.id:
        return doc, None

    member = (
        db.query(WorkspaceMember)
        .filter(
            WorkspaceMember.workspace_id == doc_id,
            WorkspaceMember.user_id == user.id,
        )
        .first()
    )

    if not member:
        raise HTTPException(403, "No tienes acceso a este workspace")

    role_hierarchy = {"viewer": 0, "editor": 1, "owner": 2}
    if role_hierarchy.get(member.role, 0) < role_hierarchy.get(required_role, 0):
        raise HTTPException(403, "No tienes permisos suficientes")

    return doc, member


def _workspace_to_dict(
    doc: WorkspaceDocument,
    members: Optional[list] = None,
    expose_share_token: bool = False,
) -> dict:
    """Serializa WorkspaceDocument a dict JSON.

    El share_link_token solo se expone cuando `expose_share_token=True`
    (solo owners). Viewers/editors no lo ven para evitar que compartan
    enlaces públicos sin autorización del owner.
    """
    result = {
        "id": doc.id,
        "title": doc.title,
        "ownerId": doc.owner_id,
        "courseName": doc.course_name,
        "apaEdition": doc.apa_edition,
        "options": doc.options,
        "isCompleted": doc.is_completed,
        "contentYjs": doc.content_yjs,
        "createdAt": doc.created_at.isoformat() if doc.created_at else "",
        "updatedAt": doc.updated_at.isoformat() if doc.updated_at else "",
        "members": members or [],
    }
    if expose_share_token:
        result["shareLinkToken"] = doc.share_link_token
    return result


def _member_to_dict(m: WorkspaceMember) -> dict:
    """Serializa WorkspaceMember a dict JSON."""
    user = m.user if hasattr(m, "user") and m.user else None
    return {
        "id": m.id,
        "workspaceId": m.workspace_id,
        "userId": m.user_id,
        "user": _user_brief(user),
        "role": m.role,
        "charsContributed": m.chars_contributed,
        "invitedAt": m.invited_at.isoformat() if m.invited_at else "",
        "joinedAt": m.joined_at.isoformat() if m.joined_at else "",
    }


def _version_to_dict(v: WorkspaceVersion) -> dict:
    """Serializa WorkspaceVersion a dict JSON."""
    return {
        "id": v.id,
        "workspaceId": v.workspace_id,
        "contentYjs": v.content_yjs,
        "createdBy": v.created_by,
        "createdAt": v.created_at.isoformat() if v.created_at else "",
        "label": v.label,
    }


def _message_to_dict(m: WorkspaceMessage) -> dict:
    """Serializa WorkspaceMessage a dict JSON.

    Patrón consistente con _member_to_dict y _version_to_dict.
    """
    return {
        "id": m.id,
        "workspaceId": m.workspace_id,
        "userId": m.user_id,
        "content": m.content,
        "createdAt": m.created_at.isoformat() if m.created_at else "",
    }


# ─── Pydantic schemas ─────────────────────────────────────────────
#
# Config:
# - populate_by_name=True permite enviar los campos en snake_case (course_name)
#   o camelCase (courseName). El frontend envía snake_case; backend los mapea.
# - alias_generator convierte automáticamente snake_case ↔ camelCase.

from pydantic import ConfigDict  # noqa: E402


def _to_camel(snake: str) -> str:
    parts = snake.split("_")
    return parts[0] + "".join(w.capitalize() for w in parts[1:])


class _WorkspaceBaseRequest(BaseModel):
    model_config = ConfigDict(
        alias_generator=_to_camel,
        populate_by_name=True,
    )


class CreateWorkspaceRequest(_WorkspaceBaseRequest):
    title: str = Field(..., min_length=1, max_length=255)
    course_name: Optional[str] = Field(None, max_length=255)
    professor: Optional[str] = Field(None, max_length=255)
    description: Optional[str] = None
    apa_edition: str = Field("7", max_length=10)
    template_id: Optional[str] = None
    options: Optional[str] = None  # JSON string
    initial_rubric_raw: Optional[str] = None


class PatchWorkspaceRequest(_WorkspaceBaseRequest):
    title: Optional[str] = Field(None, min_length=1, max_length=255)
    course_name: Optional[str] = Field(None, max_length=255)
    apa_edition: Optional[str] = Field(None, max_length=10)
    options: Optional[str] = None
    cover_data: Optional[str] = None
    is_completed: Optional[bool] = None
    # content_yjs: snapshot Yjs base64 persistido desde el cliente (sub-bloque 2b).
    # Pass-through opaco — el servidor no decodifica ni valida el contenido binario.
    content_yjs: Optional[str] = None


class AddMemberRequest(BaseModel):
    email: str = Field(..., min_length=1)
    role: str = Field("editor")


class PatchMemberRequest(BaseModel):
    role: str


class CreateVersionRequest(_WorkspaceBaseRequest):
    content_yjs: str
    label: Optional[str] = Field(None, max_length=100)


class CreateMessageRequest(BaseModel):
    # min_length no se usa aquí para retornar 400 (no 422) en mensajes vacíos.
    # La validación de contenido vacío se hace en el handler.
    content: str = Field(..., max_length=4000)


class ContributionDeltaRequest(BaseModel):
    delta_chars: int = Field(..., ge=0)


class _CitationItem(BaseModel):
    id: str
    raw: str


class ValidateCitationsRequest(BaseModel):
    citations: list[_CitationItem]


class RubricTextRequest(BaseModel):
    text: str


class CreateCommentRequest(BaseModel):
    content: str = Field(..., min_length=1, max_length=2000)
    anchor_id: str = Field(..., min_length=1)
    parent_id: Optional[str] = None


class PatchCommentRequest(BaseModel):
    content: str = Field(..., min_length=1, max_length=2000)


class ResolveCommentRequest(BaseModel):
    resolved: bool


# ─── CRUD de workspaces ───────────────────────────────────────────


@router.post("", status_code=201)
def create_workspace(
    data: CreateWorkspaceRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Crea un nuevo workspace. El creador se convierte en owner automáticamente."""
    title = data.title.strip()
    if not title:
        raise HTTPException(400, "El título es requerido")

    doc = WorkspaceDocument(
        title=title,
        owner_id=user.id,
        course_name=data.course_name,
        apa_edition=data.apa_edition or "7",
        options=data.options or "{}",
        rubric_raw=data.initial_rubric_raw,
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )
    db.add(doc)
    db.flush()  # Obtener id antes de crear el miembro

    # El owner se agrega como miembro con rol 'owner'
    owner_member = WorkspaceMember(
        workspace_id=doc.id,
        user_id=user.id,
        role="owner",
        invited_at=datetime.utcnow(),
        joined_at=datetime.utcnow(),
    )
    db.add(owner_member)
    db.commit()
    db.refresh(doc)

    members = db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == doc.id).all()
    return _workspace_to_dict(
        doc,
        [_member_to_dict(m) for m in members],
        expose_share_token=True,  # el creador es owner, puede ver el token
    )


@router.get("")
def list_workspaces(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Lista todos los workspaces donde el usuario es owner o miembro.

    Response: ``{"workspaces": [...]}`` envuelto para permitir metadata
    futura (paginación, total, etc.) sin breaking change.
    """
    owned_ids = [r[0] for r in db.query(WorkspaceDocument.id).filter(WorkspaceDocument.owner_id == user.id).all()]
    member_ids = [r[0] for r in db.query(WorkspaceMember.workspace_id).filter(WorkspaceMember.user_id == user.id).all()]

    all_ids = list(set(owned_ids + member_ids))
    if not all_ids:
        return {"workspaces": []}

    docs = (
        db.query(WorkspaceDocument)
        .filter(WorkspaceDocument.id.in_(all_ids))
        .order_by(WorkspaceDocument.updated_at.desc())
        .all()
    )

    result = []
    for doc in docs:
        members = db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == doc.id).all()
        is_owner = doc.owner_id == user.id
        d = _workspace_to_dict(
            doc,
            [_member_to_dict(m) for m in members],
            expose_share_token=is_owner,
        )
        d.pop("contentYjs", None)  # No enviar contenido Yjs en lista
        result.append(d)

    return {"workspaces": result}


@router.get("/invite/{token}")
def get_invite(
    token: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Valida un token de invitación y devuelve metadata del workspace.

    Endpoint scaffold en 2a. La generación de tokens con expiración y
    roles configurables se implementa en 2d.

    Response alineado con ``InviteTokenInfo`` del frontend.
    """
    doc = db.query(WorkspaceDocument).filter(WorkspaceDocument.share_link_token == token).first()
    if not doc:
        return {
            "valid": False,
            "workspace_id": None,
            "workspace_title": None,
            "owner_name": None,
            "proposed_role": None,
        }

    owner = db.query(User).filter(User.id == doc.owner_id).first()
    owner_name = owner.username if owner else "Desconocido"
    if owner and hasattr(owner, "first_name") and owner.first_name:
        owner_name = f"{owner.first_name} {owner.last_name or ''}".strip()

    return {
        "valid": True,
        "workspace_id": doc.id,
        "workspace_title": doc.title,
        "course_name": doc.course_name,
        "owner_name": owner_name,
        "proposed_role": "editor",  # Default scaffold 2a; rol configurable en 2d
        "token": token,
    }


@router.post("/invite/{token}/accept")
def accept_invite(
    token: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Acepta una invitación y agrega al usuario como miembro del workspace.

    Scaffold 2a: usa el ``share_link_token`` del workspace y agrega al usuario
    autenticado con rol "editor" (default). La verificación de expiración,
    roles específicos por invitación y flujo de una-sola-vez queda para 2d.
    """
    doc = db.query(WorkspaceDocument).filter(WorkspaceDocument.share_link_token == token).first()
    if not doc:
        raise HTTPException(404, "Invitación inválida o expirada")

    # Evitar duplicar miembros
    existing = (
        db.query(WorkspaceMember)
        .filter(
            WorkspaceMember.workspace_id == doc.id,
            WorkspaceMember.user_id == user.id,
        )
        .first()
    )
    if existing:
        return {"ok": True, "workspace_id": doc.id, "role": existing.role, "already_member": True}

    # Crear como editor por defecto (scaffold)
    new_member = WorkspaceMember(
        id=gen_id(),
        workspace_id=doc.id,
        user_id=user.id,
        role="editor",
        invited_at=datetime.utcnow(),
        joined_at=datetime.utcnow(),
    )
    db.add(new_member)
    db.commit()
    return {"ok": True, "workspace_id": doc.id, "role": "editor", "already_member": False}


@router.get("/{doc_id}")
def get_workspace(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Obtiene un workspace por ID. Requiere ser owner o miembro.

    El ``shareLinkToken`` solo se expone si el solicitante es owner,
    para evitar que editors/viewers compartan el enlace público sin
    autorización del owner.
    """
    doc, _ = _check_access(doc_id, user, db)
    members = db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == doc.id).all()
    is_owner = doc.owner_id == user.id
    return _workspace_to_dict(
        doc,
        [_member_to_dict(m) for m in members],
        expose_share_token=is_owner,
    )


@router.patch("/{doc_id}")
def patch_workspace(
    doc_id: str,
    data: PatchWorkspaceRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Actualiza campos del workspace. Requiere rol editor o owner."""
    doc, _ = _check_access(doc_id, user, db, required_role="editor")

    if data.title is not None:
        stripped = data.title.strip()
        if stripped:
            doc.title = stripped
    if data.course_name is not None:
        doc.course_name = data.course_name
    if data.apa_edition is not None:
        doc.apa_edition = data.apa_edition
    if data.options is not None:
        doc.options = data.options
    if data.cover_data is not None:
        doc.cover_data = data.cover_data
    if data.is_completed is not None:
        doc.is_completed = data.is_completed
    if data.content_yjs is not None:
        # Snapshot Yjs base64 — pass-through, sin validación del contenido binario.
        doc.content_yjs = data.content_yjs

    doc.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(doc)

    members = db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == doc.id).all()
    is_owner = doc.owner_id == user.id
    return _workspace_to_dict(
        doc,
        [_member_to_dict(m) for m in members],
        expose_share_token=is_owner,
    )


@router.delete("/{doc_id}")
def delete_workspace(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Elimina un workspace. Solo el owner puede hacerlo."""
    doc = db.query(WorkspaceDocument).filter(WorkspaceDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Workspace no encontrado")
    if doc.owner_id != user.id:
        raise HTTPException(403, "Solo el creador puede eliminar el workspace")

    db.delete(doc)
    db.commit()
    return {"ok": True}


# ─── Miembros ─────────────────────────────────────────────────────


@router.get("/{doc_id}/members")
def list_members(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Lista los miembros del workspace.

    Response: ``{"members": [...]}`` envuelto para permitir metadata
    futura (roles count, invitaciones pendientes) sin breaking change.
    """
    _check_access(doc_id, user, db)
    members = db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == doc_id).all()
    return {"members": [_member_to_dict(m) for m in members]}


@router.post("/{doc_id}/members", status_code=201)
def add_member(
    doc_id: str,
    data: AddMemberRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Agrega un miembro al workspace por email. Solo owner puede agregar."""
    _check_access(doc_id, user, db, required_role="owner")

    target = db.query(User).filter(User.email == data.email).first()
    if not target:
        raise HTTPException(404, "Usuario no encontrado con ese email")

    existing = (
        db.query(WorkspaceMember)
        .filter(
            WorkspaceMember.workspace_id == doc_id,
            WorkspaceMember.user_id == target.id,
        )
        .first()
    )
    if existing:
        raise HTTPException(409, "El usuario ya es miembro de este workspace")

    role = data.role if data.role in ("editor", "viewer") else "editor"

    member = WorkspaceMember(
        workspace_id=doc_id,
        user_id=target.id,
        role=role,
        invited_at=datetime.utcnow(),
    )
    db.add(member)
    db.commit()
    db.refresh(member)
    return _member_to_dict(member)


@router.patch("/{doc_id}/members/{member_id}")
def patch_member(
    doc_id: str,
    member_id: str,
    data: PatchMemberRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Cambia el rol de un miembro. Solo el owner puede hacerlo."""
    _check_access(doc_id, user, db, required_role="owner")

    member = (
        db.query(WorkspaceMember)
        .filter(
            WorkspaceMember.id == member_id,
            WorkspaceMember.workspace_id == doc_id,
        )
        .first()
    )
    if not member:
        raise HTTPException(404, "Miembro no encontrado")

    new_role = data.role if data.role in ("editor", "viewer") else "editor"
    member.role = new_role
    db.commit()
    db.refresh(member)
    return _member_to_dict(member)


@router.delete("/{doc_id}/members/{member_id}")
def remove_member(
    doc_id: str,
    member_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Remueve un miembro del workspace. Solo el owner puede remover a otros."""
    doc = db.query(WorkspaceDocument).filter(WorkspaceDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Workspace no encontrado")
    if doc.owner_id != user.id:
        raise HTTPException(403, "Solo el owner puede remover miembros")

    member = (
        db.query(WorkspaceMember)
        .filter(
            WorkspaceMember.id == member_id,
            WorkspaceMember.workspace_id == doc_id,
        )
        .first()
    )
    if not member:
        raise HTTPException(404, "Miembro no encontrado")

    # No se puede remover al owner
    if member.user_id == doc.owner_id:
        raise HTTPException(400, "No se puede remover al owner del workspace")

    db.delete(member)
    db.commit()
    return {"ok": True}


# ─── Versiones ────────────────────────────────────────────────────


@router.get("/{doc_id}/versions")
def list_versions(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Lista las versiones del workspace. Máximo 50 más recientes.

    Response: ``{"versions": [...]}`` envuelto para permitir metadata
    futura (total, cursor de paginación) sin breaking change.
    """
    _check_access(doc_id, user, db)
    versions = (
        db.query(WorkspaceVersion)
        .filter(WorkspaceVersion.workspace_id == doc_id)
        .order_by(WorkspaceVersion.created_at.desc())
        .limit(50)
        .all()
    )
    return {"versions": [_version_to_dict(v) for v in versions]}


@router.post("/{doc_id}/versions", status_code=201)
def create_version(
    doc_id: str,
    data: CreateVersionRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Crea un snapshot manual del workspace. Requiere rol editor o owner."""
    doc, _ = _check_access(doc_id, user, db, required_role="editor")

    version = WorkspaceVersion(
        workspace_id=doc_id,
        content_yjs=data.content_yjs,
        created_by=user.id,
        label=data.label,
        created_at=datetime.utcnow(),
    )
    db.add(version)
    db.commit()
    db.refresh(version)
    return _version_to_dict(version)


@router.post("/{doc_id}/versions/{version_id}/restore")
def restore_version(
    doc_id: str,
    version_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Restaura el workspace a una versión anterior. Requiere rol editor o owner."""
    doc, _ = _check_access(doc_id, user, db, required_role="editor")

    version = (
        db.query(WorkspaceVersion)
        .filter(
            WorkspaceVersion.id == version_id,
            WorkspaceVersion.workspace_id == doc_id,
        )
        .first()
    )
    if not version:
        raise HTTPException(404, "Versión no encontrada")

    # Restaurar contenido al snapshot de la versión
    doc.content_yjs = version.content_yjs
    doc.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(doc)

    members = db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == doc.id).all()
    return _workspace_to_dict(doc, [_member_to_dict(m) for m in members])


# ─── Chat grupal ──────────────────────────────────────────────────
#
# Sub-bloque 2b. Historia completa visible para todos los miembros
# (decisión §1.2.1 #2 — no filtrar por joined_at).


@router.get("/{doc_id}/chat/messages")
def list_chat_messages(
    doc_id: str,
    limit: int = 50,
    before: Optional[str] = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Lista los mensajes del chat grupal del workspace.

    Paginación: ``limit`` (máx 50) y ``before`` (ID del mensaje, retorna
    los anteriores a ese mensaje en orden descendente de creación).
    Historia completa visible para todos los miembros (decisión 2b §1.2.1 #2).
    """
    _check_access(doc_id, user, db)

    query = db.query(WorkspaceMessage).filter(WorkspaceMessage.workspace_id == doc_id)

    if before:
        pivot = db.query(WorkspaceMessage).filter(WorkspaceMessage.id == before).first()
        if pivot and pivot.created_at:
            query = query.filter(WorkspaceMessage.created_at < pivot.created_at)

    messages = query.order_by(WorkspaceMessage.created_at.desc()).limit(min(limit, 50)).all()
    # Retornar en orden cronológico ascendente para la UI
    return {"messages": [_message_to_dict(m) for m in reversed(messages)]}


@router.post("/{doc_id}/chat/messages", status_code=201)
def create_chat_message(
    doc_id: str,
    data: CreateMessageRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Envía un mensaje al chat grupal del workspace.

    Requiere ser owner o miembro (cualquier rol). Valida que el contenido
    no esté vacío (ya validado por Pydantic min_length=1, pero también
    rejected por strip).
    """
    _check_access(doc_id, user, db)

    content = data.content.strip()
    if not content:
        raise HTTPException(400, "El contenido del mensaje no puede estar vacío")

    msg = WorkspaceMessage(
        workspace_id=doc_id,
        user_id=user.id,
        content=content,
        created_at=datetime.utcnow(),
    )
    db.add(msg)
    db.commit()
    db.refresh(msg)

    logger.info(
        "Chat message created: doc=%s user=%s msg=%s",
        doc_id,
        user.id,
        msg.id,
    )
    return _message_to_dict(msg)


@router.delete("/{doc_id}/chat/messages/{msg_id}")
def delete_chat_message(
    doc_id: str,
    msg_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Elimina un mensaje del chat grupal.

    Puede eliminar:
    - El autor del mensaje.
    - El owner del workspace.
    Cualquier otro miembro recibe 403.
    """
    # Verificar acceso al workspace (cualquier rol)
    doc, _ = _check_access(doc_id, user, db)

    msg = (
        db.query(WorkspaceMessage)
        .filter(
            WorkspaceMessage.id == msg_id,
            WorkspaceMessage.workspace_id == doc_id,
        )
        .first()
    )
    if not msg:
        raise HTTPException(404, "Mensaje no encontrado")

    is_msg_author = msg.user_id == user.id
    is_workspace_owner = doc.owner_id == user.id

    if not is_msg_author and not is_workspace_owner:
        raise HTTPException(403, "Solo el autor del mensaje o el owner del workspace pueden eliminarlo")

    db.delete(msg)
    db.commit()
    return {"ok": True}


# ─── Contribution metric ──────────────────────────────────────────
#
# Sub-bloque 2b. PATCH periódico desde cliente cada 30s (decisión D6).


@router.patch("/{doc_id}/members/{member_id}/contribution")
def patch_member_contribution(
    doc_id: str,
    member_id: str,
    data: ContributionDeltaRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Incrementa el contador de caracteres contribuidos por un miembro.

    Solo el usuario dueño del ``member_id`` puede actualizar su propio
    contador (validación user.id == member.user_id, 403 en caso contrario).
    ``delta_chars`` debe ser no negativo; 0 se acepta como no-op válido.
    """
    _check_access(doc_id, user, db)

    member = (
        db.query(WorkspaceMember)
        .filter(
            WorkspaceMember.id == member_id,
            WorkspaceMember.workspace_id == doc_id,
        )
        .first()
    )
    if not member:
        raise HTTPException(404, "Miembro no encontrado")

    if member.user_id != user.id:
        raise HTTPException(403, "Solo puedes actualizar tu propia contribución")

    member.chars_contributed = (member.chars_contributed or 0) + data.delta_chars
    db.commit()
    db.refresh(member)

    return {
        "id": member.id,
        "charsContributed": member.chars_contributed,
    }


# ─── Citas APA (sub-bloque 2d.1) ──────────────────────────────────


@router.post("/{doc_id}/citations/validate")
def validate_citations(
    doc_id: str,
    data: ValidateCitationsRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Valida un batch de citas APA 7 en texto libre.

    Retorna para cada item `{id, valid, errors, suggested}`.
    Las citas con errores NO rechazan la request — el endpoint devuelve 200
    con detalle por item.
    """
    from apa_validator import validate_citation

    _check_access(doc_id, user, db, required_role="viewer")

    if not data.citations:
        raise HTTPException(400, "El campo 'citations' no puede estar vacío")

    results = []
    for item in data.citations:
        validation = validate_citation(item.raw)
        results.append(
            {
                "id": item.id,
                "valid": validation["valid"],
                "errors": validation["errors"],
                "suggested": validation["suggested"],
            }
        )

    return {"results": results}


# ─── Rúbrica (sub-bloque 2d.6) ────────────────────────────────────


@router.post("/{doc_id}/rubric/text")
def upload_rubric_text(
    doc_id: str,
    data: RubricTextRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Parsea texto plano de rúbrica y lo persiste en el doc."""
    import json as _json

    from rubric_parser import parse_rubric

    doc, _ = _check_access(doc_id, user, db, required_role="editor")

    text = (data.text or "").strip()
    if not text:
        raise HTTPException(400, "El campo 'text' no puede estar vacío")

    items, warnings = parse_rubric(text)
    doc.rubric_raw = _json.dumps({"raw": text, "items": items}, ensure_ascii=False)
    db.commit()

    return {"items": items, "warnings": warnings}


def _extract_text_from_upload(file: UploadFile) -> str:
    """Extrae texto plano de un archivo subido (PDF, DOCX, o TXT).

    Enforza `MAX_RUBRIC_UPLOAD_BYTES` leyendo en streaming; levanta
    HTTPException 413 si excede. No valida MIME ni extensión aquí
    (responsabilidad del endpoint; ver `upload_rubric_file`).
    """
    import io

    filename = (file.filename or "").lower()

    chunks: list[bytes] = []
    total = 0
    while True:
        chunk = file.file.read(1024 * 1024)  # 1 MB por vez
        if not chunk:
            break
        total += len(chunk)
        if total > MAX_RUBRIC_UPLOAD_BYTES:
            raise HTTPException(
                413,
                f"Archivo excede el tamaño máximo permitido ({MAX_RUBRIC_UPLOAD_BYTES // (1024 * 1024)} MB)",
            )
        chunks.append(chunk)
    content_bytes = b"".join(chunks)

    if filename.endswith(".pdf"):
        import pdfplumber

        parts = []
        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                parts.append(txt)
        return "\n".join(parts)

    if filename.endswith((".docx", ".doc")):
        import docx

        doc = docx.Document(io.BytesIO(content_bytes))
        return "\n".join(p.text for p in doc.paragraphs)

    # TXT u otro: decodificar como utf-8 con fallback
    try:
        return content_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return content_bytes.decode("latin-1", errors="replace")


@router.post("/{doc_id}/rubric/upload")
async def upload_rubric_file(
    doc_id: str,
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Sube archivo PDF/DOCX/TXT, extrae texto y parsea como rúbrica."""
    import json as _json

    from rubric_parser import parse_rubric

    doc, _ = _check_access(doc_id, user, db, required_role="editor")

    if not file.filename:
        raise HTTPException(400, "No se recibió archivo")

    # Validación MIME (allowlist) + extensión (defensa en profundidad)
    filename_lower = file.filename.lower()
    has_allowed_ext = any(filename_lower.endswith(ext) for ext in ALLOWED_RUBRIC_EXTENSIONS)
    mime_allowed = (file.content_type or "").split(";")[0].strip() in ALLOWED_RUBRIC_MIME_TYPES

    if not has_allowed_ext or not mime_allowed:
        logger.warning(
            "[rubric] Upload rechazado por MIME/extensión: file=%s mime=%s",
            file.filename,
            file.content_type,
        )
        raise HTTPException(
            415,
            "Tipo de archivo no permitido. Formatos aceptados: PDF, DOCX, DOC, TXT",
        )

    try:
        text = _extract_text_from_upload(file)
    except HTTPException:
        # Propagar HTTPException de _extract_text_from_upload (413 por tamaño)
        raise
    except Exception as exc:
        logger.warning("[rubric] Error extrayendo texto de %s: %s", file.filename, exc)
        raise HTTPException(400, f"No se pudo leer el archivo: {exc}") from exc

    if not text or not text.strip():
        raise HTTPException(400, "El archivo no contiene texto legible")

    items, warnings = parse_rubric(text)
    doc.rubric_raw = _json.dumps({"raw": text, "items": items}, ensure_ascii=False)
    db.commit()

    return {"items": items, "warnings": warnings, "filename": file.filename}


@router.get("/{doc_id}/rubric")
def get_rubric(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Retorna la rúbrica parseada del doc."""
    import json as _json

    doc, _ = _check_access(doc_id, user, db, required_role="viewer")

    if not doc.rubric_raw:
        return {"raw": "", "items": []}

    try:
        payload = _json.loads(doc.rubric_raw)
        return {
            "raw": payload.get("raw", ""),
            "items": payload.get("items", []),
        }
    except (ValueError, TypeError):
        # Si está en formato legacy (texto plano sin JSON), devolver raw con items vacíos
        return {"raw": doc.rubric_raw, "items": []}


# ─── Comentarios inline (sub-bloque 2d.8) ─────────────────────────


def _comment_to_dict(c) -> dict:  # noqa: ANN001
    """Serializa WorkspaceComment a dict JSON con mentions persistidas."""
    import json as _json_mod

    mentions: list[str] = []
    if c.mentions_json:
        try:
            parsed = _json_mod.loads(c.mentions_json)
            if isinstance(parsed, list):
                mentions = [str(x) for x in parsed]
        except (ValueError, TypeError):
            mentions = []

    return {
        "id": c.id,
        "workspace_id": c.workspace_id,
        "user_id": c.user_id,
        "anchor_id": c.anchor_json,
        "content": c.content,
        "resolved": bool(c.resolved),
        "parent_id": c.parent_id,
        "mentions": mentions,
        "created_at": c.created_at.isoformat() if c.created_at else None,
    }


def _extract_mentions(content: str, doc_id: str, db: Session) -> list[str]:
    """Extrae menciones @username del content y valida que sean miembros del workspace.

    Retorna lista de user_ids mencionados. Lanza HTTPException 400 si alguna
    mención no corresponde a un miembro.
    """
    import re as _re

    from database import WorkspaceMember as _WorkspaceMember  # type: ignore

    mention_pattern = _re.compile(r"@([a-zA-Z0-9_\-\.]+)")
    usernames = mention_pattern.findall(content)

    if not usernames:
        return []

    member_users = (
        db.query(User)
        .join(_WorkspaceMember, _WorkspaceMember.user_id == User.id)
        .filter(_WorkspaceMember.workspace_id == doc_id)
        .all()
    )
    username_to_id = {u.username: u.id for u in member_users}

    user_ids = []
    for uname in usernames:
        if uname not in username_to_id:
            raise HTTPException(400, f"Usuario @{uname} no es miembro de este workspace")
        user_ids.append(username_to_id[uname])

    return user_ids


@router.get("/{doc_id}/comments")
def list_comments(
    doc_id: str,
    anchor_id: Optional[str] = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Lista comentarios del doc. Filtro opcional por anchor_id."""
    from database import WorkspaceComment as _WorkspaceComment  # type: ignore

    _check_access(doc_id, user, db, required_role="viewer")

    query = db.query(_WorkspaceComment).filter(_WorkspaceComment.workspace_id == doc_id)
    if anchor_id:
        query = query.filter(_WorkspaceComment.anchor_json == anchor_id)

    comments = query.order_by(_WorkspaceComment.created_at.asc()).all()
    return {"comments": [_comment_to_dict(c) for c in comments]}


@router.post("/{doc_id}/comments", status_code=201)
def create_comment(
    doc_id: str,
    data: CreateCommentRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Crea un comentario (nuevo thread o reply con parent_id)."""
    from database import WorkspaceComment as _WorkspaceComment  # type: ignore
    from database import gen_id as _gen_id  # type: ignore

    _check_access(doc_id, user, db, required_role="viewer")

    # Validar parent_id si existe
    if data.parent_id:
        parent = (
            db.query(_WorkspaceComment)
            .filter(
                _WorkspaceComment.id == data.parent_id,
                _WorkspaceComment.workspace_id == doc_id,
            )
            .first()
        )
        if not parent:
            raise HTTPException(400, "parent_id no corresponde a un comentario de este workspace")

    # Extraer y validar menciones (lanza 400 si usuario no-miembro mencionado)
    import json as _json_mod

    mentions = _extract_mentions(data.content, doc_id, db)

    comment = _WorkspaceComment(
        id=_gen_id(),
        workspace_id=doc_id,
        user_id=user.id,
        anchor_json=data.anchor_id,
        content=data.content,
        parent_id=data.parent_id,
        mentions_json=_json_mod.dumps(mentions) if mentions else None,
    )
    db.add(comment)
    db.commit()
    db.refresh(comment)

    return _comment_to_dict(comment)


@router.patch("/{doc_id}/comments/{comment_id}")
def patch_comment(
    doc_id: str,
    comment_id: str,
    data: PatchCommentRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Edita contenido de un comentario. Solo autor puede editar."""
    from database import WorkspaceComment as _WorkspaceComment  # type: ignore

    _check_access(doc_id, user, db, required_role="viewer")

    comment = (
        db.query(_WorkspaceComment)
        .filter(
            _WorkspaceComment.id == comment_id,
            _WorkspaceComment.workspace_id == doc_id,
        )
        .first()
    )
    if not comment:
        raise HTTPException(404, "Comentario no encontrado")

    if comment.user_id != user.id:
        raise HTTPException(403, "Solo el autor puede editar su comentario")

    # Re-validar menciones del nuevo content y persistir
    import json as _json_mod

    mentions = _extract_mentions(data.content, doc_id, db)

    comment.content = data.content
    comment.mentions_json = _json_mod.dumps(mentions) if mentions else None
    db.commit()
    db.refresh(comment)
    return _comment_to_dict(comment)


@router.delete("/{doc_id}/comments/{comment_id}")
def delete_comment(
    doc_id: str,
    comment_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Borra comentario. Permitido para autor o owner del workspace."""
    from database import WorkspaceComment as _WorkspaceComment  # type: ignore
    from database import WorkspaceDocument as _WD  # type: ignore

    _check_access(doc_id, user, db, required_role="viewer")

    comment = (
        db.query(_WorkspaceComment)
        .filter(
            _WorkspaceComment.id == comment_id,
            _WorkspaceComment.workspace_id == doc_id,
        )
        .first()
    )
    if not comment:
        raise HTTPException(404, "Comentario no encontrado")

    doc = db.query(_WD).filter(_WD.id == doc_id).first()
    is_owner = doc and doc.owner_id == user.id

    if comment.user_id != user.id and not is_owner:
        raise HTTPException(403, "Solo el autor o el owner pueden borrar este comentario")

    db.delete(comment)
    db.commit()
    return {"ok": True}


@router.post("/{doc_id}/comments/{comment_id}/resolve")
def resolve_comment(
    doc_id: str,
    comment_id: str,
    data: ResolveCommentRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    """Marca thread como resuelto/no-resuelto. Solo owner o editor."""
    from database import WorkspaceComment as _WorkspaceComment  # type: ignore

    _check_access(doc_id, user, db, required_role="editor")

    comment = (
        db.query(_WorkspaceComment)
        .filter(
            _WorkspaceComment.id == comment_id,
            _WorkspaceComment.workspace_id == doc_id,
        )
        .first()
    )
    if not comment:
        raise HTTPException(404, "Comentario no encontrado")

    comment.resolved = data.resolved
    db.commit()
    db.refresh(comment)
    return _comment_to_dict(comment)
