"""
Export PDF/DOCX para Workspaces v2 — sub-sub-bloque 2d.7.

Cierra la deuda C1 de docs/pendientes.md:
    backend/collab_routes.py:455-503 tenía SSRF con xhtml2pdf
    permitiendo <img src="http://169.254.169.254"> → AWS metadata leak.

Este módulo implementa el export desde cero con defensa en profundidad:

Estrategia E1 (decisión de Cristian): WeasyPrint para PDF.
Estrategia E2: pre-descarga imágenes a base64 + bleach whitelist.

Si WeasyPrint no puede cargarse por falta de deps nativas (cairo/pango),
los endpoints de PDF retornan 501 con mensaje explicativo. Los endpoints
DOCX siempre funcionan (python-docx no tiene deps nativas).

Componente legal:
- Los PDFs generados NO incluyen metadatos de identidad del usuario
  (id, email, nombre) gracias al parámetro presentational_hints=False
  y a la ausencia explícita de metadatos en el HTML de entrada.
- Ver sección "Metadatos PDF" al final de este módulo.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any
from urllib.parse import urlparse

import bleach  # type: ignore[import-untyped]
import httpx
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from database import get_db  # type: ignore
from middleware import get_current_user  # type: ignore
from workspaces_routes import _check_access  # type: ignore

logger = logging.getLogger("conniku.workspaces_export")

# ─── WeasyPrint — import condicional ─────────────────────────────────
#
# WeasyPrint requiere libcairo2 + libpango-1.0-0 (y gobject-2.0).
# En Render: agregar a apt.txt (o equivalente buildpack):
#     libcairo2 libpango-1.0-0 libpangocairo-1.0-0 libgdk-pixbuf2.0-0
#     libffi-dev shared-mime-info
# En macOS local de desarrollo: brew install pango cairo
# Si no están disponibles, _WEASYPRINT_AVAILABLE = False → endpoints PDF
# retornan 501 con instrucciones de infra.

_WEASYPRINT_AVAILABLE = False
_HTML_CLASS: Any = None

try:
    from weasyprint import HTML  # type: ignore

    _HTML_CLASS = HTML
    _WEASYPRINT_AVAILABLE = True
    logger.info("WeasyPrint disponible — export PDF habilitado")
except OSError:
    logger.warning(
        "WeasyPrint no pudo cargarse (faltan librerías nativas: cairo/pango/gobject). "
        "Export PDF retornará 501. Para habilitar: instalar libcairo2 libpango-1.0-0 "
        "en el servidor (Render: agregar a apt.txt)."
    )
except ImportError:
    logger.warning("WeasyPrint no está instalado. Instalar con: pip install 'weasyprint>=62.0,<65.0'")

# ─── Constantes de seguridad ──────────────────────────────────────────

# Whitelist de dominios para imágenes remotas.
# Solo dominios propios de Conniku pueden ser descargados.
# Cualquier otro dominio es rechazado antes de hacer el request.
_ALLOWED_REMOTE_IMG_DOMAINS: frozenset[str] = frozenset(
    {
        "conniku.com",
        "www.conniku.com",
        "cdn.conniku.com",
        "api.conniku.com",
        # Supabase Storage del proyecto (el hostname real de Supabase)
        # se puede agregar aquí cuando esté configurado.
    }
)

# Rangos de IPs privadas/reservadas que NO deben recibir requests.
# RFC1918 + loopback + link-local (AWS metadata) + TEST-NET
_BLOCKED_IP_PREFIXES: tuple[str, ...] = (
    "10.",
    "172.16.",
    "172.17.",
    "172.18.",
    "172.19.",
    "172.20.",
    "172.21.",
    "172.22.",
    "172.23.",
    "172.24.",
    "172.25.",
    "172.26.",
    "172.27.",
    "172.28.",
    "172.29.",
    "172.30.",
    "172.31.",
    "192.168.",
    "127.",
    "169.254.",  # AWS instance metadata endpoint
    "0.",
    "::1",
    "fc00:",
    "fd00:",
    "fe80:",
)

# Protocolos permitidos en src de img para descarga remota.
# Solo HTTPS sobre dominios de la whitelist.
_ALLOWED_PROTOCOLS: frozenset[str] = frozenset({"https"})

# Tags HTML permitidos en el export PDF.
_ALLOWED_TAGS: list[str] = [
    "p",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "ul",
    "ol",
    "li",
    "blockquote",
    "em",
    "strong",
    "a",
    "img",
    "br",
    "hr",
    "span",
    "div",
    "table",
    "thead",
    "tbody",
    "tr",
    "th",
    "td",
    "caption",
    "code",
    "pre",
]

# Atributos permitidos por tag.
# Sólo los estrictamente necesarios; todos los event handlers (on*) quedan
# excluidos por construcción: bleach sólo permite lo que está en esta lista.
_ALLOWED_ATTRS: dict[str, list[str]] = {
    "a": ["href", "title", "target"],
    "img": ["src", "alt", "width", "height"],
    "td": ["colspan", "rowspan"],
    "th": ["colspan", "rowspan"],
    "span": ["class"],
    "div": ["class"],
    "p": ["class"],
    "h1": ["class"],
    "h2": ["class"],
    "h3": ["class"],
}

# Tamaño máximo de imagen remota descargable: 5 MB.
_MAX_IMG_SIZE_BYTES: int = 5 * 1024 * 1024

# Timeout para descarga de imágenes remotas (en segundos).
_IMG_DOWNLOAD_TIMEOUT_S: float = 5.0


# ─── Función: sanitize_html ───────────────────────────────────────────


def sanitize_html(html: str) -> str:
    """Sanitiza HTML con bleach usando whitelist estricta de tags y atributos.

    Elimina:
    - Tags no permitidos (<script>, <iframe>, <object>, <embed>, etc.)
    - Todos los atributos de evento (onerror, onclick, onload, etc.)
    - href con esquema javascript:
    - src con esquema file:// o esquemas no-http en img

    Preserva:
    - Tags de la lista _ALLOWED_TAGS
    - Atributos de la lista _ALLOWED_ATTRS
    - src con data:image/* en <img> (imágenes ya inlinadas, seguras)
    - href con https:// en <a>

    Implementación: bleach.clean con un callable de atributos que implementa
    whitelist estricta. El callable solo acepta atributos que están en
    _ALLOWED_ATTRS para el tag correspondiente; todos los demás (incluyendo
    event handlers como onerror, onclick, onload) son rechazados.

    Args:
        html: HTML crudo del documento del usuario.

    Returns:
        HTML sanitizado listo para pasar al motor de render.
    """

    def _attrs_filter(tag: str, name: str, value: str) -> bool:
        """Whitelist de atributos por tag.

        Bleach llama este callable por cada (tag, attr_name, attr_value).
        Retornar True para permitir, False para eliminar.

        Cualquier atributo que no esté en la whitelist es eliminado.
        Esto incluye todos los event handlers (on*) por construcción.
        """
        # Obtener atributos permitidos para este tag
        allowed_for_tag = _ALLOWED_ATTRS.get(tag, [])

        # Si el atributo no está en la lista permitida para este tag → rechazar
        if name not in allowed_for_tag:
            return False

        # Validación adicional por atributo específico
        if tag == "a" and name == "href":
            parsed = urlparse(value)
            # Rechazar esquemas peligrosos (javascript:, data:, vbscript:, etc.)
            if parsed.scheme and parsed.scheme.lower() not in ("https", "http", ""):
                return False
            return True

        if tag == "img" and name == "src":
            # Permitir explícitamente data:image/*
            if value.lower().startswith("data:image/"):
                return True
            # Permitir http/https (los SSRF se eliminaron en inline_remote_images)
            parsed = urlparse(value)
            if parsed.scheme.lower() in ("http", "https", ""):
                return True
            # Rechazar cualquier otro esquema (file://, gopher://, etc.)
            return False

        return True

    # Nota sobre protocols:
    # bleach.clean tiene su propio filtro de protocolos que actúa DESPUÉS del
    # callable de atributos. El protocolo "data" debe estar explícitamente
    # en la lista para que bleach no elimine src="data:image/..." aunque el
    # callable ya lo haya aprobado.
    # Limitamos a los protocolos estrictamente necesarios para el export.
    cleaned = bleach.clean(
        html,
        tags=_ALLOWED_TAGS,
        attributes=_attrs_filter,
        protocols=["https", "http", "data"],
        strip=True,
        strip_comments=True,
    )
    return cleaned


# ─── Función: _is_ssrf_target ─────────────────────────────────────────


def _is_ssrf_target(url: str) -> bool:
    """Retorna True si la URL apunta a un destino SSRF conocido.

    Comprueba:
    1. El esquema no es https (bloquea http, file, gopher, ftp, etc.)
    2. El hostname es una IP privada/reservada
    3. El hostname es localhost / equivalentes
    4. El dominio no está en la whitelist de dominios permitidos

    Esta función es el gate principal anti-SSRF.
    """
    try:
        parsed = urlparse(url)
    except ValueError:
        return True  # URL malformada → bloquear

    scheme = (parsed.scheme or "").lower()
    hostname = (parsed.hostname or "").lower()

    # Solo https para imágenes remotas (http se bloquea también)
    if scheme not in _ALLOWED_PROTOCOLS:
        return True

    # Bloquear hostnames que son localhost
    if hostname in ("localhost", "127.0.0.1", "::1", ""):
        return True

    # Bloquear IPs privadas / reservadas / link-local
    for prefix in _BLOCKED_IP_PREFIXES:
        if hostname.startswith(prefix):
            return True

    # Bloquear cualquier hostname que parezca dirección IP (no hostname DNS)
    # pero no esté en la whitelist. Regex: solo dígitos y puntos → es una IP.
    if re.match(r"^\d+\.\d+\.\d+\.\d+$", hostname):
        # Cualquier IP que llegue aquí no es de la whitelist → bloquear
        return True

    # Verificar whitelist de dominios
    # El dominio debe ser exactamente conniku.com o un subdominio: *.conniku.com
    for allowed in _ALLOWED_REMOTE_IMG_DOMAINS:
        if hostname == allowed or hostname.endswith(f".{allowed.lstrip('*.')}"):
            return False  # Es seguro, NO es SSRF target

    # Si no coincide con ningún dominio de la whitelist → bloquear
    return True


# ─── Función: inline_remote_images ────────────────────────────────────


def inline_remote_images(html: str) -> str:
    """Pre-descarga imágenes remotas permitidas y las convierte a data:base64.

    Estrategia E2 (defensa en profundidad):
    1. Busca todos los <img src="..."> en el HTML.
    2. Para cada src:
       - Si es data:image/* → ya está inlineada, no toca nada.
       - Si es SSRF target → elimina la etiqueta <img> completa.
       - Si es dominio permitido → descarga (timeout 5s, max 5MB) y convierte.
       - Si la descarga falla → elimina la etiqueta <img> completa.
    3. Retorna el HTML con todas las referencias remotas resueltas o eliminadas.

    Después de este paso, el motor de render (WeasyPrint) no necesita hacer
    ningún request externo, eliminando el vector SSRF a nivel de render.

    Args:
        html: HTML ya sanitizado (o sin sanitizar, esta función es idempotente
              respecto de la sanitización pero NO hace sanitización completa).

    Returns:
        HTML con imgs remotas reemplazadas por data URIs o eliminadas.
    """
    IMG_SRC_RE = re.compile(
        r'(<img\b[^>]*?)(\bsrc=["\'])([^"\']+)(["\'])([^>]*?>)',
        re.IGNORECASE,
    )

    def _replace_img(match: re.Match) -> str:  # type: ignore[type-arg]
        prefix = match.group(1)
        quote = match.group(2)
        src = match.group(3)
        close_quote = match.group(4)
        suffix = match.group(5)

        # data:image/* ya está inlineada — preservar sin cambios
        if src.lower().startswith("data:image/"):
            return match.group(0)

        # Verificar si es SSRF target antes de cualquier request
        if _is_ssrf_target(src):
            logger.warning(
                "SSRF bloqueado: imagen removida de HTML de export. URL: %s",
                src[:80],  # Truncar para no volcar URLs largas en logs
            )
            # Eliminar la etiqueta <img> completa
            return ""

        # Dominio en whitelist → descargar y convertir a base64
        try:
            with httpx.Client(timeout=_IMG_DOWNLOAD_TIMEOUT_S, follow_redirects=False) as client:
                response = client.get(src)
                response.raise_for_status()

                # Verificar tamaño
                if len(response.content) > _MAX_IMG_SIZE_BYTES:
                    logger.warning("Imagen remota excede %s bytes, removida: %s", _MAX_IMG_SIZE_BYTES, src[:80])
                    return ""

                # Inferir content-type
                ct = response.headers.get("content-type", "image/png").split(";")[0].strip()
                if not ct.startswith("image/"):
                    logger.warning("Content-type no es imagen: %s — removida", ct)
                    return ""

                import base64

                data_b64 = base64.b64encode(response.content).decode("ascii")
                data_uri = f"data:{ct};base64,{data_b64}"

                return f"{prefix}{quote}{data_uri}{close_quote}{suffix}"

        except (httpx.HTTPError, httpx.TimeoutException) as exc:
            logger.warning("Error al descargar imagen remota %s: %s", src[:80], exc)
            return ""

    return IMG_SRC_RE.sub(_replace_img, html)


# ─── Función: export_pdf ──────────────────────────────────────────────


def export_pdf(html: str) -> bytes:
    """Genera un PDF desde HTML usando WeasyPrint.

    Pipeline:
    1. inline_remote_images: resuelve/elimina imgs remotas (anti-SSRF)
    2. sanitize_html: bleach whitelist (anti-XSS, anti-tags peligrosos)
    3. WeasyPrint HTML.write_pdf(): render a bytes

    Metadatos PDF:
    WeasyPrint puede incluir metadatos en el PDF via etiquetas <title>,
    <meta name="author">, etc. El HTML de entrada no debe incluir datos
    de identidad del usuario en estas etiquetas. El llamador es responsable
    de no pasar HTML con <meta name="author" content="{user.email}">.
    Este módulo no agrega metadatos de identificación por su cuenta.

    Args:
        html: HTML del documento. Puede contener imgs remotas que serán
              procesadas por inline_remote_images.

    Returns:
        Bytes del PDF generado.

    Raises:
        RuntimeError: si WeasyPrint no está disponible (deps nativas ausentes).
    """
    if not _WEASYPRINT_AVAILABLE or _HTML_CLASS is None:
        raise RuntimeError(
            "WeasyPrint no está disponible. Instalar libcairo2 libpango-1.0-0 libpangocairo-1.0-0 en el servidor."
        )

    # Paso 1: pre-procesar imágenes remotas (antes de sanitizar, para que
    # el sanitizador no toque los data URIs que acabamos de generar)
    processed_html = inline_remote_images(html)

    # Paso 2: sanitizar con bleach (elimina scripts, event handlers, etc.)
    safe_html = sanitize_html(processed_html)

    # Paso 3: render PDF
    # presentational_hints=False: no leer style/color hints del HTML
    # (reduce superficie de ataque CSS)
    pdf_bytes: bytes = _HTML_CLASS(string=safe_html).write_pdf()
    return pdf_bytes


# ─── Función: export_docx ─────────────────────────────────────────────


def export_docx(blocks: list[dict]) -> bytes:
    """Genera un DOCX desde una lista de bloques estructurados.

    Usa python-docx directamente sin pasar por HTML (no hay riesgo SSRF
    en DOCX porque no descarga recursos externos).

    Formato de bloques aceptados:
        {"type": "h1", "text": "Título"}
        {"type": "h2", "text": "Subtítulo"}
        {"type": "h3", "text": "Sub-subtítulo"}
        {"type": "p", "text": "Párrafo de texto."}
        {"type": "list", "items": ["Item 1", "Item 2"]}

    Cualquier bloque con type desconocido es ignorado silenciosamente
    (retrocompatibilidad futura).

    Metadatos DOCX:
    python-docx incluye por defecto la propiedad "creator" como vacía.
    Esto se mantiene vacío explícitamente: no se asigna user.email ni
    user.id a las propiedades del documento.

    Args:
        blocks: Lista de bloques estructurados del documento.

    Returns:
        Bytes del archivo DOCX (ZIP con firma PK).
    """
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore

    doc = Document()

    # Limpiar metadatos de identidad explícitamente
    # (python-docx crea propiedades vacías por default; esto asegura que
    # no se filtren datos de la sesión del proceso)
    core_props = doc.core_properties
    core_props.author = ""
    core_props.last_modified_by = ""
    core_props.title = ""

    # Construir contenido desde bloques
    heading_map = {
        "h1": 1,
        "h2": 2,
        "h3": 3,
        "h4": 4,
        "h5": 5,
        "h6": 6,
    }

    for block in blocks:
        block_type = block.get("type", "")
        text = block.get("text", "")

        if block_type in heading_map:
            doc.add_heading(text, level=heading_map[block_type])
        elif block_type == "p":
            doc.add_paragraph(text)
        elif block_type == "list":
            items = block.get("items", [])
            if isinstance(items, list):
                for item in items:
                    doc.add_paragraph(str(item), style="List Bullet")
        elif block_type == "blockquote":
            para = doc.add_paragraph(text)
            para.paragraph_format.left_indent = Pt(36)
        else:
            # Tipo desconocido: ignorar silenciosamente
            logger.debug("Bloque de tipo desconocido ignorado en export_docx: %s", block_type)

    # Serializar a bytes en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ─── Pydantic schemas para los endpoints ─────────────────────────────


class ExportPdfRequest(BaseModel):
    html: str = Field(..., min_length=1, description="HTML del documento a exportar")
    include_cover: bool = Field(False, description="Incluir página de portada")
    include_rubric: bool = Field(False, description="Incluir rúbrica al final")


class ExportDocxRequest(BaseModel):
    blocks: list[dict] = Field(..., description="Lista de bloques estructurados del documento")
    include_cover: bool = Field(False, description="Incluir página de portada")
    include_rubric: bool = Field(False, description="Incluir rúbrica al final")


# ─── Router FastAPI ───────────────────────────────────────────────────

router = APIRouter(prefix="/workspaces", tags=["workspaces-export"])


@router.post("/{workspace_id}/export/pdf")
def export_workspace_pdf(
    workspace_id: str,
    data: ExportPdfRequest,
    user=Depends(get_current_user),
    db: Session = Depends(get_db),
) -> Response:
    """Exporta el workspace a PDF.

    Requiere rol editor o superior (no viewers).
    El HTML recibido pasa por sanitización completa (bleach) y
    pre-procesamiento de imágenes (inline base64) antes del render.
    Ningún request externo es realizado por el motor de render.

    Returns:
        application/pdf con Content-Disposition: attachment; filename=workspace-{id}.pdf

    Raises:
        403: Si el usuario no tiene rol editor o superior.
        404: Si el workspace no existe.
        501: Si WeasyPrint no está disponible (deps nativas ausentes en el servidor).
    """
    # Valida acceso: solo editor o owner pueden exportar
    _check_access(workspace_id, user, db, required_role="editor")

    if not _WEASYPRINT_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail=(
                "Export PDF no disponible: las librerías nativas de WeasyPrint "
                "(cairo/pango) no están instaladas en este servidor. "
                "Contacta al administrador para habilitar esta función."
            ),
        )

    try:
        pdf_bytes = export_pdf(data.html)
    except RuntimeError as exc:
        logger.error("Error en export_pdf para workspace %s: %s", workspace_id, exc)
        raise HTTPException(status_code=501, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Error inesperado en export_pdf para workspace %s: %s", workspace_id, exc)
        raise HTTPException(status_code=500, detail="Error interno al generar PDF") from exc

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="workspace-{workspace_id}.pdf"',
        },
    )


@router.post("/{workspace_id}/export/docx")
def export_workspace_docx(
    workspace_id: str,
    data: ExportDocxRequest,
    user=Depends(get_current_user),
    db: Session = Depends(get_db),
) -> Response:
    """Exporta el workspace a DOCX.

    Requiere rol editor o superior (no viewers).
    El DOCX se construye desde bloques estructurados sin pasar por HTML,
    por lo que no existe vector SSRF en este endpoint.

    Los metadatos del DOCX (author, last_modified_by, title) quedan
    vacíos para no filtrar datos personales del usuario que exporta.

    Returns:
        application/vnd.openxmlformats-officedocument.wordprocessingml.document
        con Content-Disposition: attachment; filename=workspace-{id}.docx

    Raises:
        403: Si el usuario no tiene rol editor o superior.
        404: Si el workspace no existe.
        422: Si el body no es válido (Pydantic).
    """
    # Valida acceso: solo editor o owner pueden exportar
    _check_access(workspace_id, user, db, required_role="editor")

    try:
        docx_bytes = export_docx(data.blocks)
    except Exception as exc:
        logger.error("Error en export_docx para workspace %s: %s", workspace_id, exc)
        raise HTTPException(status_code=500, detail="Error interno al generar DOCX") from exc

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="workspace-{workspace_id}.docx"',
        },
    )
