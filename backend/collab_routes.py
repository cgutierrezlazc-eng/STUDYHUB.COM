"""
Collaborative Documents — Trabajos Grupales.

Real-time collaborative document editing for group assignments.
"""
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, Query
from sqlalchemy.orm import Session
from sqlalchemy import desc, or_

from database import (
    get_db, User, gen_id,
    CollabDocument, CollabDocumentMember, CollabDocumentVersion,
    CollabDocumentMessage,
)
from middleware import get_current_user

router = APIRouter(prefix="/collab", tags=["collab"])


def _user_brief(u):
    if not u:
        return None
    return {
        "id": u.id, "username": u.username, "firstName": u.first_name,
        "lastName": u.last_name, "avatar": u.avatar,
    }


def _doc_to_dict(doc: CollabDocument, members: list | None = None) -> dict:
    return {
        "id": doc.id,
        "title": doc.title,
        "description": doc.description,
        "content": doc.content,
        "ownerId": doc.owner_id,
        "owner": _user_brief(doc.owner),
        "status": doc.status,
        "university": doc.university,
        "career": doc.career,
        "courseName": doc.course_name,
        "color": doc.color,
        "icon": doc.icon,
        "members": members or [],
        "createdAt": doc.created_at.isoformat() if doc.created_at else "",
        "updatedAt": doc.updated_at.isoformat() if doc.updated_at else "",
    }


def _member_to_dict(m: CollabDocumentMember) -> dict:
    return {
        "id": m.id,
        "userId": m.user_id,
        "user": _user_brief(m.user),
        "role": m.role,
        "joinedAt": m.joined_at.isoformat() if m.joined_at else "",
    }


def _check_access(doc_id: str, user: User, db: Session, min_role: str = "viewer"):
    """Check user has access to doc. Returns (doc, member_record)."""
    doc = db.query(CollabDocument).filter(CollabDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Documento no encontrado")

    # Owner always has full access
    if doc.owner_id == user.id:
        return doc, None

    member = db.query(CollabDocumentMember).filter(
        CollabDocumentMember.document_id == doc_id,
        CollabDocumentMember.user_id == user.id,
    ).first()

    if not member:
        raise HTTPException(403, "No tienes acceso a este documento")

    role_hierarchy = {"viewer": 0, "editor": 1, "owner": 2}
    if role_hierarchy.get(member.role, 0) < role_hierarchy.get(min_role, 0):
        raise HTTPException(403, "No tienes permisos suficientes")

    return doc, member


# ─── CRUD ────────────────────────────────────────────────────────

@router.post("")
def create_document(data: dict, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    title = data.get("title", "").strip()
    if not title:
        raise HTTPException(400, "Titulo requerido")

    doc = CollabDocument(
        id=gen_id(),
        title=title,
        description=data.get("description", ""),
        owner_id=user.id,
        university=data.get("university", user.university or ""),
        career=data.get("career", user.career or ""),
        course_name=data.get("courseName", ""),
        color=data.get("color", "#2D62C8"),
        icon=data.get("icon", "file-text"),
    )
    db.add(doc)

    # Owner is automatically a member with 'owner' role
    owner_member = CollabDocumentMember(
        id=gen_id(), document_id=doc.id, user_id=user.id, role="owner",
    )
    db.add(owner_member)

    # Add initial members if provided
    member_ids = data.get("memberIds", [])
    for uid in member_ids:
        if uid == user.id:
            continue
        target = db.query(User).filter(User.id == uid).first()
        if target:
            m = CollabDocumentMember(
                id=gen_id(), document_id=doc.id, user_id=uid, role="editor",
            )
            db.add(m)

    db.commit()
    db.refresh(doc)

    members = db.query(CollabDocumentMember).filter(CollabDocumentMember.document_id == doc.id).all()
    return _doc_to_dict(doc, [_member_to_dict(m) for m in members])


@router.get("")
def list_documents(
    status: str = Query("active"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    # Documents where user is owner OR member
    owned_ids = db.query(CollabDocument.id).filter(
        CollabDocument.owner_id == user.id, CollabDocument.status == status
    ).all()
    member_ids = db.query(CollabDocumentMember.document_id).filter(
        CollabDocumentMember.user_id == user.id
    ).all()

    all_ids = list(set([r[0] for r in owned_ids] + [r[0] for r in member_ids]))
    if not all_ids:
        return []

    docs = db.query(CollabDocument).filter(
        CollabDocument.id.in_(all_ids), CollabDocument.status == status
    ).order_by(desc(CollabDocument.updated_at)).all()

    result = []
    for doc in docs:
        members = db.query(CollabDocumentMember).filter(
            CollabDocumentMember.document_id == doc.id
        ).all()
        d = _doc_to_dict(doc, [_member_to_dict(m) for m in members])
        d.pop("content", None)  # Don't send content in list view
        result.append(d)

    return result


@router.get("/{doc_id}")
def get_document(doc_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc, _ = _check_access(doc_id, user, db)
    members = db.query(CollabDocumentMember).filter(
        CollabDocumentMember.document_id == doc.id
    ).all()
    return _doc_to_dict(doc, [_member_to_dict(m) for m in members])


@router.put("/{doc_id}")
def update_document(doc_id: str, data: dict, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc, _ = _check_access(doc_id, user, db, min_role="editor")

    if "title" in data:
        doc.title = data["title"].strip() or doc.title
    if "description" in data:
        doc.description = data["description"]
    if "content" in data:
        doc.content = data["content"]
    if "color" in data:
        doc.color = data["color"]
    if "icon" in data:
        doc.icon = data["icon"]
    if "courseName" in data:
        doc.course_name = data["courseName"]
    if "status" in data and doc.owner_id == user.id:
        doc.status = data["status"]

    doc.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(doc)

    members = db.query(CollabDocumentMember).filter(
        CollabDocumentMember.document_id == doc.id
    ).all()
    return _doc_to_dict(doc, [_member_to_dict(m) for m in members])


@router.delete("/{doc_id}")
def delete_document(doc_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(CollabDocument).filter(CollabDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Documento no encontrado")
    if doc.owner_id != user.id:
        raise HTTPException(403, "Solo el creador puede eliminar el documento")

    db.delete(doc)
    db.commit()
    return {"ok": True}


# ─── Members ─────────────────────────────────────────────────────

@router.post("/{doc_id}/members")
def add_member(doc_id: str, data: dict, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc, _ = _check_access(doc_id, user, db, min_role="editor")

    user_id = data.get("userId", "").strip()
    if not user_id:
        raise HTTPException(400, "userId requerido")

    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        raise HTTPException(404, "Usuario no encontrado")

    existing = db.query(CollabDocumentMember).filter(
        CollabDocumentMember.document_id == doc_id,
        CollabDocumentMember.user_id == user_id,
    ).first()
    if existing:
        raise HTTPException(400, "El usuario ya es miembro")

    role = data.get("role", "editor")
    if role not in ("editor", "viewer"):
        role = "editor"

    m = CollabDocumentMember(
        id=gen_id(), document_id=doc_id, user_id=user_id, role=role,
    )
    db.add(m)
    db.commit()
    db.refresh(m)
    return _member_to_dict(m)


@router.delete("/{doc_id}/members/{member_user_id}")
def remove_member(doc_id: str, member_user_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(CollabDocument).filter(CollabDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Documento no encontrado")

    # Only owner can remove others; members can remove themselves
    if user.id != doc.owner_id and user.id != member_user_id:
        raise HTTPException(403, "No tienes permisos para eliminar miembros")

    if member_user_id == doc.owner_id:
        raise HTTPException(400, "No puedes eliminar al creador del documento")

    member = db.query(CollabDocumentMember).filter(
        CollabDocumentMember.document_id == doc_id,
        CollabDocumentMember.user_id == member_user_id,
    ).first()
    if not member:
        raise HTTPException(404, "Miembro no encontrado")

    db.delete(member)
    db.commit()
    return {"ok": True}


@router.put("/{doc_id}/members/{member_user_id}/role")
def update_member_role(doc_id: str, member_user_id: str, data: dict, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(CollabDocument).filter(CollabDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Documento no encontrado")
    if doc.owner_id != user.id:
        raise HTTPException(403, "Solo el creador puede cambiar roles")

    member = db.query(CollabDocumentMember).filter(
        CollabDocumentMember.document_id == doc_id,
        CollabDocumentMember.user_id == member_user_id,
    ).first()
    if not member:
        raise HTTPException(404, "Miembro no encontrado")

    role = data.get("role", "editor")
    if role not in ("editor", "viewer"):
        role = "editor"

    member.role = role
    db.commit()
    return _member_to_dict(member)


# ─── Versions ────────────────────────────────────────────────────

@router.get("/{doc_id}/versions")
def list_versions(doc_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    _check_access(doc_id, user, db)
    versions = db.query(CollabDocumentVersion).filter(
        CollabDocumentVersion.document_id == doc_id
    ).order_by(desc(CollabDocumentVersion.version_number)).limit(50).all()

    return [{
        "id": v.id,
        "versionNumber": v.version_number,
        "createdBy": _user_brief(v.author),
        "createdAt": v.created_at.isoformat() if v.created_at else "",
    } for v in versions]


@router.post("/{doc_id}/versions")
def save_version(doc_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc, _ = _check_access(doc_id, user, db, min_role="editor")

    last = db.query(CollabDocumentVersion).filter(
        CollabDocumentVersion.document_id == doc_id
    ).order_by(desc(CollabDocumentVersion.version_number)).first()

    version_number = (last.version_number + 1) if last else 1

    v = CollabDocumentVersion(
        id=gen_id(),
        document_id=doc_id,
        content=doc.content or "",
        version_number=version_number,
        created_by=user.id,
    )
    db.add(v)
    db.commit()

    return {
        "id": v.id,
        "versionNumber": v.version_number,
        "createdBy": _user_brief(user),
        "createdAt": v.created_at.isoformat() if v.created_at else "",
    }


@router.get("/{doc_id}/versions/{version_id}")
def get_version(doc_id: str, version_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    _check_access(doc_id, user, db)
    v = db.query(CollabDocumentVersion).filter(
        CollabDocumentVersion.id == version_id,
        CollabDocumentVersion.document_id == doc_id,
    ).first()
    if not v:
        raise HTTPException(404, "Version no encontrada")

    return {
        "id": v.id,
        "versionNumber": v.version_number,
        "content": v.content,
        "createdBy": _user_brief(v.author),
        "createdAt": v.created_at.isoformat() if v.created_at else "",
    }


# ─── User search (for adding members) ────────────────────────────

@router.get("/users/search")
def search_users(
    q: str = Query("", min_length=1),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Search users by name/username/email for adding to documents."""
    query = q.strip()
    if len(query) < 2:
        return []

    results = db.query(User).filter(
        User.id != user.id,
        or_(
            User.username.ilike(f"%{query}%"),
            User.first_name.ilike(f"%{query}%"),
            User.last_name.ilike(f"%{query}%"),
            User.email.ilike(f"%{query}%"),
        )
    ).limit(10).all()

    return [_user_brief(u) for u in results]


# ─── Document Chat ───────────────────────────────────────────────

@router.get("/{doc_id}/chat")
def get_chat_messages(
    doc_id: str,
    before: str = Query(None),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _check_access(doc_id, user, db)
    q = db.query(CollabDocumentMessage).filter(CollabDocumentMessage.document_id == doc_id)
    if before:
        q = q.filter(CollabDocumentMessage.created_at < before)
    messages = q.order_by(desc(CollabDocumentMessage.created_at)).limit(50).all()
    messages.reverse()

    return [{
        "id": m.id,
        "userId": m.user_id,
        "user": _user_brief(m.user),
        "content": m.content,
        "createdAt": m.created_at.isoformat() if m.created_at else "",
    } for m in messages]


@router.post("/{doc_id}/chat")
async def send_chat_message(
    doc_id: str,
    data: dict,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _check_access(doc_id, user, db)
    content = data.get("content", "").strip()
    if not content:
        raise HTTPException(400, "Mensaje vacio")

    msg = CollabDocumentMessage(
        id=gen_id(), document_id=doc_id, user_id=user.id, content=content,
    )
    db.add(msg)
    db.commit()

    msg_data = {
        "id": msg.id,
        "userId": msg.user_id,
        "user": _user_brief(user),
        "content": msg.content,
        "createdAt": msg.created_at.isoformat() if msg.created_at else "",
    }

    # Broadcast via document WebSocket room
    try:
        from collab_ws import _rooms
        room = _rooms.get(doc_id)
        if room:
            import asyncio
            await room.broadcast_json({"type": "chat_message", "message": msg_data})
    except Exception:
        pass  # WS broadcast is best-effort

    return msg_data


# ─── Export (PDF / DOCX) ─────────────────────────────────────────

@router.get("/{doc_id}/export/pdf")
def export_pdf(doc_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc, _ = _check_access(doc_id, user, db)
    raw_html = doc.content or "<p>Documento vacio</p>"

    # Fix C1 SSRF (2026-04-19 Cristian aprobó opción b — parchear V1 directo):
    # aplicar la misma defensa en profundidad que workspaces_export.py (V2):
    # 1. bleach.clean() remueve scripts + event handlers + iframe
    # 2. inline_remote_images pre-descarga imágenes de dominios whitelisted
    #    (conniku.com/cdn/api), rechaza todo el resto (incluido 169.254.169.254
    #    AWS metadata, file://, gopher://, IPs RFC1918). Imágenes fuera de
    #    whitelist se eliminan silenciosamente del HTML antes del render.
    # xhtml2pdf.pisa.CreatePDF recibe HTML ya sanitizado + sin remote fetches.
    from workspaces_export import inline_remote_images, sanitize_html

    sanitized = sanitize_html(raw_html)
    html_content = inline_remote_images(sanitized)

    import io
    from xhtml2pdf import pisa

    # Wrap in a full HTML page with styling
    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: Helvetica, Arial, sans-serif; font-size: 12px; line-height: 1.6; color: #1a1a1a; margin: 40px; }}
  h1 {{ font-size: 24px; font-weight: bold; margin: 20px 0 10px; }}
  h2 {{ font-size: 20px; font-weight: bold; margin: 18px 0 8px; }}
  h3 {{ font-size: 16px; font-weight: bold; margin: 14px 0 6px; }}
  p {{ margin: 0 0 8px; }}
  ul, ol {{ padding-left: 24px; margin: 8px 0; }}
  blockquote {{ border-left: 3px solid #2D62C8; padding-left: 12px; color: #555; font-style: italic; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
  th, td {{ border: 1px solid #ddd; padding: 6px 10px; text-align: left; }}
  th {{ background: #f5f5f5; font-weight: bold; }}
  code {{ background: #f0f0f0; padding: 1px 4px; border-radius: 3px; font-size: 11px; }}
  pre {{ background: #f0f0f0; padding: 10px; border-radius: 4px; font-size: 11px; overflow-x: auto; }}
  mark {{ background: #FBBF24; padding: 0 2px; }}
  img {{ max-width: 100%; }}
</style>
</head><body>
<h1 style="text-align:center; margin-bottom:4px;">{doc.title}</h1>
<p style="text-align:center; color:#888; font-size:10px; margin-bottom:30px;">
  {doc.course_name or ''} &mdash; Conniku
</p>
<hr style="border:none; border-top:1px solid #ddd; margin-bottom:20px;">
{html_content}
</body></html>"""

    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(full_html, dest=buffer, encoding='utf-8')
    if pisa_status.err:
        raise HTTPException(500, "Error generando PDF")

    buffer.seek(0)
    from fastapi.responses import StreamingResponse
    safe_title = "".join(c for c in doc.title if c.isalnum() or c in " -_").strip()[:60]
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
    )


@router.get("/{doc_id}/export/docx")
def export_docx(doc_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc, _ = _check_access(doc_id, user, db)
    html_content = doc.content or ""

    import io
    import re
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = DocxDocument()

    # Title
    title_para = document.add_heading(doc.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if doc.course_name:
        sub = document.add_paragraph(doc.course_name)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        sub.runs[0].font.size = Pt(10)

    document.add_paragraph("")  # Spacer

    # Simple HTML to DOCX conversion
    # Strip tags and process basic elements
    text = html_content
    # Convert <br> to newlines
    text = re.sub(r'<br\s*/?>', '\n', text)
    # Convert headers
    for level in [3, 2, 1]:
        pattern = rf'<h{level}[^>]*>(.*?)</h{level}>'
        for match in re.finditer(pattern, text, re.DOTALL):
            heading_text = re.sub(r'<[^>]+>', '', match.group(1)).strip()
            if heading_text:
                document.add_heading(heading_text, level=level)
        text = re.sub(pattern, '', text, flags=re.DOTALL)

    # Convert lists
    for match in re.finditer(r'<li[^>]*>(.*?)</li>', text, re.DOTALL):
        item_text = re.sub(r'<[^>]+>', '', match.group(1)).strip()
        if item_text:
            document.add_paragraph(item_text, style='List Bullet')
    text = re.sub(r'</?[uo]l[^>]*>', '', text)
    text = re.sub(r'<li[^>]*>.*?</li>', '', text, flags=re.DOTALL)

    # Convert paragraphs and remaining text
    text = re.sub(r'</?(?:div|section|article|span)[^>]*>', '', text)
    paragraphs = re.split(r'</?p[^>]*>', text)
    for p_text in paragraphs:
        clean = re.sub(r'<[^>]+>', '', p_text).strip()
        if clean:
            para = document.add_paragraph()
            # Handle bold/italic inline
            parts = re.split(r'(<(?:strong|b|em|i|u)>.*?</(?:strong|b|em|i|u)>)', p_text)
            for part in parts:
                clean_part = re.sub(r'<[^>]+>', '', part).strip()
                if not clean_part:
                    continue
                run = para.add_run(clean_part)
                if '<strong>' in part or '<b>' in part:
                    run.bold = True
                if '<em>' in part or '<i>' in part:
                    run.italic = True
                if '<u>' in part:
                    run.underline = True

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    from fastapi.responses import StreamingResponse
    safe_title = "".join(c for c in doc.title if c.isalnum() or c in " -_").strip()[:60]
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
    )
